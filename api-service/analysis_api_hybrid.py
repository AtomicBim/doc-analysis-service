"""
FastAPI —Å–µ—Ä–≤–∏—Å –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏ —Å –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–µ–º –≥–∏–±—Ä–∏–¥–Ω–æ–≥–æ –ø–æ–¥—Ö–æ–¥–∞:
- –¢–ó/–¢–£: —Ä—É—á–Ω–æ–π –ø–∞—Ä—Å–∏–Ω–≥ –∏ —Å–µ–≥–º–µ–Ω—Ç–∞—Ü–∏—è —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π
- –ß–µ—Ä—Ç–µ–∂–∏: OpenAI Assistants API —Å File Search –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞
"""
import os
import json
import logging
import asyncio
import warnings
from typing import List, Optional, Dict, Any, Tuple
from pathlib import Path
from tenacity import retry, stop_after_attempt, wait_exponential
from openai import RateLimitError

# –û—Ç–∫–ª—é—á–∞–µ–º warnings –æ deprecation
warnings.filterwarnings("ignore", category=DeprecationWarning, module="openai")

import uvicorn
from fastapi import FastAPI, HTTPException, Form, UploadFile, File, Request
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from openai import AsyncOpenAI
from dotenv import load_dotenv

# –ò–º–ø–æ—Ä—Ç –Ω–∞—à–µ–≥–æ —Ä–µ—Ñ–∞–∫—Ç–æ—Ä–∏–Ω–≥–æ–≤–æ–≥–æ PDF –ø—Ä–æ—Ü–µ—Å—Å–æ—Ä–∞
from pdf_processor import PDFProcessor, PDFBatchProcessor
from progress_tracker import ProgressTracker
from config import (
    # Stage 1
    STAGE1_MAX_PAGES, STAGE1_DPI, STAGE1_QUALITY, STAGE1_MAX_PAGES_PER_REQUEST,
    STAGE1_STAMP_CROP, STAGE1_TOP_RIGHT_CROP, STAGE1_HEADER_CROP, STAGE1_BOTTOM_CENTER_CROP,
    # Stage 2
    STAGE2_MAX_PAGES, STAGE2_DPI, STAGE2_QUALITY, STAGE2_DETAIL, STAGE2_MAX_PAGES_PER_REQUEST,
    # Stage 3
    STAGE3_DPI, STAGE3_QUALITY, STAGE3_DETAIL, STAGE3_BATCH_SIZE, STAGE3_MAX_COMPLETION_TOKENS, STAGE3_RETRY_ON_REFUSAL, STAGE3_MAX_PAGES_PER_REQUEST,
    # Stage 4
    STAGE4_ENABLED, STAGE4_SAMPLE_PAGES_PER_SECTION, STAGE4_DPI, STAGE4_QUALITY, STAGE4_DETAIL, STAGE4_MAX_COMPLETION_TOKENS,
    # Retry
    RETRY_MAX_ATTEMPTS, RETRY_WAIT_EXPONENTIAL_MULTIPLIER, RETRY_WAIT_EXPONENTIAL_MAX,
    # OpenAI
    OPENAI_MODEL,
    # Logging
    LOG_LEVEL, LOG_RESPONSE_PREVIEW_LENGTH, LOG_FULL_RESPONSE_ON_ERROR
)

# ============================
# –ö–û–ù–§–ò–ì–£–†–ê–¶–ò–Ø
# ============================

logging.basicConfig(level=getattr(logging, LOG_LEVEL), format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# –ó–∞–≥—Ä—É–∑–∫–∞ –ø–µ—Ä–µ–º–µ–Ω–Ω—ã—Ö –æ–∫—Ä—É–∂–µ–Ω–∏—è
load_dotenv()

# –ò–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏—è OpenAI API
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    logger.error("OPENAI_API_KEY –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω –≤ –ø–µ—Ä–µ–º–µ–Ω–Ω—ã—Ö –æ–∫—Ä—É–∂–µ–Ω–∏—è!")
    raise ValueError("OPENAI_API_KEY is required")

client = AsyncOpenAI(api_key=OPENAI_API_KEY)
MAX_FILE_SIZE_MB = 40
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024

logger.info(f"üöÄ –ò—Å–ø–æ–ª—å–∑—É–µ—Ç—Å—è OpenAI API (VISION MODE): {OPENAI_MODEL}")
logger.info("üìã –ê—Ä—Ö–∏—Ç–µ–∫—Ç—É—Ä–∞: –¢–ó/–¢–£ –ø–∞—Ä—Å–∏–Ω–≥ –≤—Ä—É—á–Ω—É—é + –ß–µ—Ä—Ç–µ–∂–∏ —á–µ—Ä–µ–∑ Vision API")


# ============================
# –°–ò–°–¢–ï–ú–ê –ü–†–û–ú–ü–¢–û–í
# ============================

def load_prompts() -> Dict[str, str]:
    """–ó–∞–≥—Ä—É–∂–∞–µ—Ç –ø—Ä–æ–º–ø—Ç—ã –∏–∑ —Ñ–∞–π–ª–æ–≤ –≤ –ø–∞–ø–∫–µ prompts."""
    prompts = {}
    prompts_dir = Path(__file__).parent.parent / "prompts"

    stage_files = {
        "–ì–ö": "gk_prompt.txt",
        "–§–≠": "fe_prompt.txt",
        "–≠–ü": "ep_prompt.txt"
    }

    for stage, filename in stage_files.items():
        file_path = prompts_dir / filename
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                prompts[stage] = f.read().strip()
            logger.info(f"‚úÖ –ó–∞–≥—Ä—É–∂–µ–Ω –ø—Ä–æ–º–ø—Ç –¥–ª—è —Å—Ç–∞–¥–∏–∏ {stage}")
        except FileNotFoundError:
            logger.error(f"‚ùå –ù–µ –Ω–∞–π–¥–µ–Ω —Ñ–∞–π–ª –ø—Ä–æ–º–ø—Ç–∞: {file_path}")
            raise FileNotFoundError(f"–§–∞–π–ª –ø—Ä–æ–º–ø—Ç–∞ –Ω–µ –Ω–∞–π–¥–µ–Ω: {file_path}")

    return prompts


def load_requirements_extraction_prompt() -> str:
    """–ó–∞–≥—Ä—É–∂–∞–µ—Ç –ø—Ä–æ–º–ø—Ç –¥–ª—è –∏–∑–≤–ª–µ—á–µ–Ω–∏—è —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π –∏–∑ –¢–ó."""
    prompts_dir = Path(__file__).parent.parent / "prompts"
    file_path = prompts_dir / "requirements_extraction_prompt.txt"
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            prompt = f.read().strip()
        logger.info(f"‚úÖ –ó–∞–≥—Ä—É–∂–µ–Ω –ø—Ä–æ–º–ø—Ç –¥–ª—è –∏–∑–≤–ª–µ—á–µ–Ω–∏—è —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π")
        return prompt
    except FileNotFoundError:
        logger.error(f"‚ùå –ù–µ –Ω–∞–π–¥–µ–Ω —Ñ–∞–π–ª –ø—Ä–æ–º–ø—Ç–∞: {file_path}")
        raise FileNotFoundError(f"–§–∞–π–ª –ø—Ä–æ–º–ø—Ç–∞ –Ω–µ –Ω–∞–π–¥–µ–Ω: {file_path}")


def load_stage_prompts() -> Dict[str, str]:
    """–ó–∞–≥—Ä—É–∂–∞–µ—Ç –ø—Ä–æ–º–ø—Ç—ã –¥–ª—è –≤—Å–µ—Ö —Å—Ç–∞–¥–∏–π –∞–Ω–∞–ª–∏–∑–∞."""
    prompts_dir = Path(__file__).parent.parent / "prompts"
    stage_prompts: Dict[str, str] = {}
    
    stage_files = {
        "stage1_metadata": "stage1_metadata_extraction_prompt.txt",
        "stage2_relevance": "stage2_page_relevance_prompt.txt",
        "stage3_analysis": "stage3_detailed_analysis_prompt.txt",
        "stage4_contradictions": "stage4_contradictions_prompt.txt",
        "system_prompt_template": "analysis_system_prompt_template.txt"
    }
    
    for key, filename in stage_files.items():
        file_path = prompts_dir / filename
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                stage_prompts[key] = f.read().strip()
            logger.info(f"‚úÖ –ó–∞–≥—Ä—É–∂–µ–Ω –ø—Ä–æ–º–ø—Ç {key}")
        except FileNotFoundError:
            logger.error(f"‚ùå –ù–µ –Ω–∞–π–¥–µ–Ω —Ñ–∞–π–ª –ø—Ä–æ–º–ø—Ç–∞: {file_path}")
            raise FileNotFoundError(f"–§–∞–π–ª –ø—Ä–æ–º–ø—Ç–∞ –Ω–µ –Ω–∞–π–¥–µ–Ω: {file_path}")
    
    return stage_prompts


# –ó–∞–≥—Ä—É–∂–∞–µ–º –ø—Ä–æ–º–ø—Ç—ã –ø—Ä–∏ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∞—Ü–∏–∏
PROMPTS = load_prompts()
REQUIREMENTS_EXTRACTION_PROMPT = load_requirements_extraction_prompt()
STAGE_PROMPTS = load_stage_prompts()

# ============================
# PDF PROCESSING –§–£–ù–ö–¶–ò–ò
# ============================

def normalize_sheet_number_to_digit(sheet_num: str) -> str:
    """
    –ù–æ—Ä–º–∞–ª–∏–∑—É–µ—Ç –Ω–æ–º–µ—Ä –ª–∏—Å—Ç–∞ –∫ –¢–û–õ–¨–ö–û –¶–ò–§–†–û–í–û–ú–£ —Ñ–æ—Ä–º–∞—Ç—É.
    –ü—Ä–∏–º–µ—Ä—ã:
        "–ê–†-01" ‚Üí "1"
        "–ö–†-03.1" ‚Üí "3"
        "–õ–∏—Å—Ç 26" ‚Üí "26"
        "5" ‚Üí "5"
    """
    import re

    if not sheet_num or sheet_num == "N/A":
        return "N/A"

    # –ò–∑–≤–ª–µ–∫–∞–µ–º –≤—Å–µ —á–∏—Å–ª–∞ –∏–∑ —Å—Ç—Ä–æ–∫–∏
    numbers = re.findall(r'\d+', str(sheet_num))

    if not numbers:
        logger.warning(f"‚ö†Ô∏è –ù–µ —É–¥–∞–ª–æ—Å—å –∏–∑–≤–ª–µ—á—å —Ü–∏—Ñ—Ä—ã –∏–∑ –Ω–æ–º–µ—Ä–∞ –ª–∏—Å—Ç–∞: '{sheet_num}'")
        return "N/A"

    # –ë–µ—Ä–µ–º –ü–û–°–õ–ï–î–ù–ï–ï —á–∏—Å–ª–æ (–æ–±—ã—á–Ω–æ —ç—Ç–æ –Ω–æ–º–µ—Ä –ª–∏—Å—Ç–∞ –≤ —Ñ–æ—Ä–º–∞—Ç–∞—Ö —Ç–∏–ø–∞ "–ê–†-01")
    digit_only = numbers[-1]

    # –£–±–∏—Ä–∞–µ–º –≤–µ–¥—É—â–∏–µ –Ω—É–ª–∏: "01" ‚Üí "1"
    digit_only = str(int(digit_only))

    logger.debug(f"üìã –ù–æ—Ä–º–∞–ª–∏–∑–∞—Ü–∏—è –Ω–æ–º–µ—Ä–∞ –ª–∏—Å—Ç–∞: '{sheet_num}' ‚Üí '{digit_only}'")
    return digit_only


def _combine_crops_for_metadata(crops: List[str]) -> str:
    """
    –û–±—ä–µ–¥–∏–Ω—è–µ—Ç 4 crops (header, top_right, bottom_center, stamp) –≤ –æ–¥–Ω–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ
    –¥–ª—è –∫–æ–º–ø–∞–∫—Ç–Ω–æ–≥–æ –∞–Ω–∞–ª–∏–∑–∞ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö —Å—Ç—Ä–∞–Ω–∏—Ü—ã.
    """
    import base64
    from PIL import Image
    import io

    # –î–µ–∫–æ–¥–∏—Ä—É–µ–º base64 –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
    images = []
    for crop_b64 in crops:
        img_data = base64.b64decode(crop_b64)
        img = Image.open(io.BytesIO(img_data))
        images.append(img)

    # –ü—Ä–µ–¥–ø–æ–ª–∞–≥–∞–µ–º, —á—Ç–æ –≤—Å–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –∏–º–µ—é—Ç –æ–¥–∏–Ω–∞–∫–æ–≤—ã–π —Ä–∞–∑–º–µ—Ä
    # –í –±—É–¥—É—â–µ–º –º–æ–∂–Ω–æ —Å–¥–µ–ª–∞—Ç—å –±–æ–ª–µ–µ –≥–∏–±–∫–∏–º
    width, height = images[0].size

    # –°–æ–∑–¥–∞–µ–º –∫–æ–º–±–∏–Ω–∏—Ä–æ–≤–∞–Ω–Ω–æ–µ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ
    combined = Image.new('RGB', (width, int(height * 0.55)))

    # –†–∞–∑–º–µ—â–∞–µ–º crops:
    # - header: –≤–µ—Ä—Ö–Ω—è—è —á–∞—Å—Ç—å (0, 0)
    combined.paste(images[0], (0, 0))

    # - top_right: –ø—Ä–∞–≤—ã–π –≤–µ—Ä—Ö–Ω–∏–π —É–≥–æ–ª
    combined.paste(images[1], (int(width * 0.7), int(height * 0.1)))

    # - bottom_center: —Ü–µ–Ω—Ç—Ä –≤–Ω–∏–∑—É
    combined.paste(images[2], (int(width * 0.3), int(height * 0.2)))

    # - stamp: –ø—Ä–∞–≤—ã–π –Ω–∏–∂–Ω–∏–π —É–≥–æ–ª
    combined.paste(images[3], (int(width * 0.7), int(height * 0.3)))

    # –°–æ—Ö—Ä–∞–Ω—è–µ–º –∫–∞–∫ base64
    img_byte_arr = io.BytesIO()
    combined.save(img_byte_arr, format='JPEG', quality=STAGE1_QUALITY)
    return base64.b64encode(img_byte_arr.getvalue()).decode('utf-8')

async def extract_selected_pdf_pages_as_images(
    doc_content: bytes,
    filename: str,
    selected_pages: List[int],
    detail: str = "low",
    dpi: int = 100,
    quality: int = 70
) -> Tuple[List[str], List[int]]:
    """
    –ò–∑–≤–ª–µ–∫–∞–µ—Ç –¢–û–õ–¨–ö–û –≤—ã–±—Ä–∞–Ω–Ω—ã–µ —Å—Ç—Ä–∞–Ω–∏—Ü—ã PDF –∫–∞–∫ base64-encoded –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è.
    –í–æ–∑–≤—Ä–∞—â–∞–µ—Ç –∫–æ—Ä—Ç–µ–∂ (images, page_numbers) –≤ —Ç–æ–π –∂–µ –ø–æ—Å–ª–µ–¥–æ–≤–∞—Ç–µ–ª—å–Ω–æ—Å—Ç–∏.

    –¢–µ–ø–µ—Ä—å –∏—Å–ø–æ–ª—å–∑—É–µ—Ç –æ–ø—Ç–∏–º–∏–∑–∏—Ä–æ–≤–∞–Ω–Ω—ã–π PDFBatchProcessor –¥–ª—è –ø–∞—Ä–∞–ª–ª–µ–ª—å–Ω–æ–π –æ–±—Ä–∞–±–æ—Ç–∫–∏.
    """
    logger.info(f"üìÑ [IMG] –ò–∑–≤–ª–µ—á–µ–Ω–∏–µ –≤—ã–±—Ä–∞–Ω–Ω—ã—Ö —Å—Ç—Ä–∞–Ω–∏—Ü –∏–∑ {filename}: {selected_pages[:10]}{'...' if len(selected_pages) > 10 else ''} (detail={detail})")

    processor = PDFBatchProcessor(doc_content, filename)
    images = await processor.extract_pages_batch(selected_pages, dpi, quality)

    # –§–∏–ª—å—Ç—Ä—É–µ–º –ø—É—Å—Ç—ã–µ —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã (–Ω–µ—É–¥–∞—á–Ω—ã–µ —Å—Ç—Ä–∞–Ω–∏—Ü—ã)
    valid_images = []
    valid_page_nums = []

    for img, page_num in zip(images, selected_pages):
        if img:  # –ù–µ –ø—É—Å—Ç–∞—è —Å—Ç—Ä–æ–∫–∞
            valid_images.append(img)
            valid_page_nums.append(page_num)

    logger.info(f"‚úÖ [IMG] –ò–∑–≤–ª–µ—á–µ–Ω–æ {len(valid_images)}/{len(selected_pages)} –≤—ã–±—Ä–∞–Ω–Ω—ã—Ö —Å—Ç—Ä–∞–Ω–∏—Ü")
    return valid_images, valid_page_nums


async def extract_page_metadata(doc_content: bytes, filename: str, max_pages: int = None) -> List[Dict[str, Any]]:
    """
    Stage 1: –ò–∑–≤–ª–µ–∫–∞–µ—Ç –º–µ—Ç–∞–¥–∞–Ω–Ω—ã–µ —Å—Ç—Ä–∞–Ω–∏—Ü (—à—Ç–∞–º–ø, –∑–∞–≥–æ–ª–æ–≤–∫–∏) —á–µ—Ä–µ–∑ Vision API.
    –§–æ–∫—É—Å –Ω–∞ –ø—Ä–∞–≤—ã–π –Ω–∏–∂–Ω–∏–π —É–≥–æ–ª (—à—Ç–∞–º–ø), –ø—Ä–∞–≤—ã–π –≤–µ—Ä—Ö–Ω–∏–π, –∑–∞–≥–æ–ª–æ–≤–∫–∏.
    """
    if max_pages is None:
        max_pages = STAGE1_MAX_PAGES

    logger.info(f"üìã [STAGE 1] –ò–∑–≤–ª–µ—á–µ–Ω–∏–µ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö –∏–∑ {filename}...")

    def _extract_crops():
        # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –æ–±–ª–∞—Å—Ç–∏ –¥–ª—è –≤—ã—Ä–µ–∑–∞–Ω–∏—è
        crop_areas = [
            STAGE1_HEADER_CROP,       # –ó–∞–≥–æ–ª–æ–≤–æ–∫
            STAGE1_TOP_RIGHT_CROP,    # –ü—Ä–∞–≤—ã–π –≤–µ—Ä—Ö–Ω–∏–π —É–≥–æ–ª
            STAGE1_BOTTOM_CENTER_CROP, # –°–µ—Ä–µ–¥–∏–Ω–∞ –≤–Ω–∏–∑—É
            STAGE1_STAMP_CROP         # –®—Ç–∞–º–ø (–ø—Ä–∞–≤—ã–π –Ω–∏–∂–Ω–∏–π)
        ]

        with PDFProcessor(doc_content, filename) as processor:
            metadata_images = []
            total_pages = min(processor.page_count, max_pages)

            for page_num in range(1, total_pages + 1):  # 1-based
                # –ò–∑–≤–ª–µ–∫–∞–µ–º –≤—Å–µ crops –¥–ª—è —Å—Ç—Ä–∞–Ω–∏—Ü—ã
                crops = processor.extract_page_crops(page_num, crop_areas, STAGE1_DPI, STAGE1_QUALITY)

                if len(crops) == 4:
                    # –û–±—ä–µ–¥–∏–Ω—è–µ–º crops –≤ –æ–¥–Ω–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏–µ –¥–ª—è –∫–æ–º–ø–∞–∫—Ç–Ω–æ—Å—Ç–∏
                    # –≠—Ç–æ –¥–µ–ª–∞–µ—Ç—Å—è –≤ –æ—Ç–¥–µ–ª—å–Ω–æ–π —Ñ—É–Ω–∫—Ü–∏–∏ –¥–ª—è —á–∏—Ç–∞–µ–º–æ—Å—Ç–∏
                    combined_base64 = _combine_crops_for_metadata(crops)
                    metadata_images.append({
                        'page_number': page_num,
                        'image': combined_base64
                    })

            logger.info(f"‚úÖ [STAGE 1] –ò–∑–≤–ª–µ—á–µ–Ω–æ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö —Å {len(metadata_images)} —Å—Ç—Ä–∞–Ω–∏—Ü")
            return metadata_images

    crops = await asyncio.to_thread(_extract_crops)

    # –§—É–Ω–∫—Ü–∏—è –¥–ª—è –æ–±—Ä–∞–±–æ—Ç–∫–∏ –æ–¥–Ω–æ–≥–æ –±–∞—Ç—á–∞ —Å retry
    @retry(stop=stop_after_attempt(3), wait=wait_exponential(multiplier=1, min=2, max=10))
    async def process_batch_with_retry(batch_crops: List[Dict], batch_start: int, batch_end: int) -> List[Dict[str, Any]]:
        """–û–±—Ä–∞–±–æ—Ç–∫–∞ –±–∞—Ç—á–∞ —Å—Ç—Ä–∞–Ω–∏—Ü —Å retry –ø—Ä–∏ –æ—à–∏–±–∫–∞—Ö"""
        content = [{
            "type": "text",
            "text": STAGE_PROMPTS["stage1_metadata"]
        }]

        # –î–æ–±–∞–≤–ª—è–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –±–∞—Ç—á–∞
        for item in batch_crops:
            content.append({
                "type": "text",
                "text": f"\n--- –°—Ç—Ä–∞–Ω–∏—Ü–∞ {item['page_number']} ---"
            })
            content.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/jpeg;base64,{item['image']}",
                    "detail": "low"
                }
            })

        response = await client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[{"role": "user", "content": content}],
            response_format={"type": "json_object"},
            max_completion_tokens=4000
        )

        response_text = response.choices[0].message.content
        if not response_text or not response_text.strip():
            raise ValueError(f"Empty response from OpenAI for batch {batch_start}-{batch_end}")

        try:
            data = json.loads(response_text)
            batch_metadata = data.get('pages', [])
            logger.info(f"‚úÖ [STAGE 1] –ë–∞—Ç—á {batch_start}-{batch_end}: –∏–∑–≤–ª–µ—á–µ–Ω–æ {len(batch_metadata)} –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö")
            return batch_metadata
        except json.JSONDecodeError as e:
            logger.error(f"‚ùå [STAGE 1] JSON parse error for batch {batch_start}-{batch_end}: {e}")
            logger.error(f"Response preview: {response_text[:200]}...")
            raise

    # –ü—Ä–æ–≤–µ—Ä—è–µ–º –ª–∏–º–∏—Ç —Å—Ç—Ä–∞–Ω–∏—Ü (–∑–∞—â–∏—Ç–∞ –æ—Ç 429 rate limit)
    if len(crops) > STAGE1_MAX_PAGES_PER_REQUEST:
        logger.warning(f"‚ö†Ô∏è [STAGE 1] –°–ª–∏—à–∫–æ–º –º–Ω–æ–≥–æ —Å—Ç—Ä–∞–Ω–∏—Ü ({len(crops)} > {STAGE1_MAX_PAGES_PER_REQUEST})")
        logger.warning(f"‚ö†Ô∏è [STAGE 1] –†–∞–∑–±–∏–≤–∞–µ–º –Ω–∞ –±–∞—Ç—á–∏ –ø–æ {STAGE1_MAX_PAGES_PER_REQUEST} —Å—Ç—Ä–∞–Ω–∏—Ü...")

        all_pages_metadata = []
        for batch_start in range(0, len(crops), STAGE1_MAX_PAGES_PER_REQUEST):
            batch_end = min(batch_start + STAGE1_MAX_PAGES_PER_REQUEST, len(crops))
            batch_crops = crops[batch_start:batch_end]

            logger.info(f"üìÑ [STAGE 1] –û–±—Ä–∞–±–æ—Ç–∫–∞ –±–∞—Ç—á–∞ —Å—Ç—Ä–∞–Ω–∏—Ü {batch_start+1}-{batch_end}...")

            try:
                # –ò—Å–ø–æ–ª—å–∑—É–µ–º —Ñ—É–Ω–∫—Ü–∏—é —Å retry
                batch_metadata = await process_batch_with_retry(batch_crops, batch_start+1, batch_end)
                all_pages_metadata.extend(batch_metadata)

            except Exception as e:
                logger.error(f"‚ùå [STAGE 1] –í—Å–µ –ø–æ–ø—ã—Ç–∫–∏ –¥–ª—è –±–∞—Ç—á–∞ {batch_start+1}-{batch_end} –ø—Ä–æ–≤–∞–ª–∏–ª–∏—Å—å: {e}")
                # Fallback –¥–ª—è —ç—Ç–æ–≥–æ –±–∞—Ç—á–∞
                logger.warning(f"‚ö†Ô∏è [STAGE 1] –ò—Å–ø–æ–ª—å–∑—É–µ–º fallback –¥–ª—è –±–∞—Ç—á–∞ {batch_start+1}-{batch_end}")
                for item in batch_crops:
                    all_pages_metadata.append({
                        "page": item['page_number'],
                        "title": f"–°—Ç—Ä–∞–Ω–∏—Ü–∞ {item['page_number']}",
                        "section": "Unknown",
                        "type": "unknown",
                        "sheet_number": f"{item['page_number']}",
                        "sheet_number_validation": {
                            "matches": False,
                            "found_in": ["stamp", "bottom_center", "top_right"],
                            "values": ["N/A", "N/A", "N/A"]
                        }
                    })

        logger.info(f"‚úÖ [STAGE 1] –ò–∑–≤–ª–µ—á–µ–Ω–æ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö –¥–ª—è {len(all_pages_metadata)} —Å—Ç—Ä–∞–Ω–∏—Ü (–±–∞—Ç—á–∏ –ø–æ {STAGE1_MAX_PAGES_PER_REQUEST})")
        return all_pages_metadata

    # –ï—Å–ª–∏ —Å—Ç—Ä–∞–Ω–∏—Ü –º–∞–ª–æ - –æ—Ç–ø—Ä–∞–≤–ª—è–µ–º –æ–¥–Ω–∏–º –∑–∞–ø—Ä–æ—Å–æ–º (–æ—Ä–∏–≥–∏–Ω–∞–ª—å–Ω–∞—è –ª–æ–≥–∏–∫–∞)
    logger.info(f"üîç [STAGE 1] –ê–Ω–∞–ª–∏–∑ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö —á–µ—Ä–µ–∑ Vision API...")

    try:
        # –ò—Å–ø–æ–ª—å–∑—É–µ–º —Ñ—É–Ω–∫—Ü–∏—é —Å retry –¥–ª—è –æ–¥–∏–Ω–æ—á–Ω–æ–≥–æ –∑–∞–ø—Ä–æ—Å–∞
        pages_metadata = await process_batch_with_retry(crops, 1, len(crops))
        logger.info(f"‚úÖ [STAGE 1] –ò–∑–≤–ª–µ—á–µ–Ω–æ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö –¥–ª—è {len(pages_metadata)} —Å—Ç—Ä–∞–Ω–∏—Ü")
        return pages_metadata

    except Exception as e:
        logger.error(f"‚ùå [STAGE 1] –í—Å–µ –ø–æ–ø—ã—Ç–∫–∏ –∏–∑–≤–ª–µ—á–µ–Ω–∏—è –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö –ø—Ä–æ–≤–∞–ª–∏–ª–∏—Å—å: {e}")
        # Fallback: –≤–æ–∑–≤—Ä–∞—â–∞–µ–º –ø—É—Å—Ç—ã–µ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã–µ
        logger.warning(f"‚ö†Ô∏è [STAGE 1] –ò—Å–ø–æ–ª—å–∑—É–µ–º fallback –¥–ª—è –≤—Å–µ—Ö {len(crops)} —Å—Ç—Ä–∞–Ω–∏—Ü")
        return [{
            "page": i+1,
            "title": f"–°—Ç—Ä–∞–Ω–∏—Ü–∞ {i+1}",
            "section": "Unknown",
            "type": "unknown",
            "sheet_number": f"{i+1}",
            "sheet_number_validation": {
                "matches": False,
                "found_in": ["stamp", "bottom_center", "top_right"],
                "values": ["N/A", "N/A", "N/A"]
            }
        } for i in range(len(crops))]


async def assess_page_relevance(
    pages_metadata: List[Dict[str, Any]],
    doc_images_low: List[str],
    requirements: List[Dict[str, Any]],
    page_numbers: Optional[List[int]] = None
) -> Dict[int, List[int]]:
    """
    Stage 2: –û—Ü–µ–Ω–∫–∞ —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω–æ—Å—Ç–∏ —Å—Ç—Ä–∞–Ω–∏—Ü –¥–ª—è –∫–∞–∂–¥–æ–≥–æ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è.
    –í–æ–∑–≤—Ä–∞—â–∞–µ—Ç mapping: {requirement_number: [page_numbers]}

    –û–ø—Ç–∏–º–∏–∑–∏—Ä–æ–≤–∞–Ω–æ –¥–ª—è gpt-5-mini: –≤—Å–µ–≥–¥–∞ –∏—Å–ø–æ–ª—å–∑—É–µ–º Vision API —Å high-res
    """
    logger.info(f"üîç [STAGE 2] –û—Ü–µ–Ω–∫–∞ —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω–æ—Å—Ç–∏ {len(pages_metadata)} —Å—Ç—Ä–∞–Ω–∏—Ü –¥–ª—è {len(requirements)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π...")

    # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –Ω—É–∂–Ω–æ –ª–∏ —Ä–∞–∑–±–∏–≤–∞—Ç—å –Ω–∞ –±–∞—Ç—á–∏
    if len(doc_images_low) > STAGE2_MAX_PAGES_PER_REQUEST:
        logger.warning(f"‚ö†Ô∏è [STAGE 2] –ú–Ω–æ–≥–æ —Å—Ç—Ä–∞–Ω–∏—Ü ({len(doc_images_low)} > {STAGE2_MAX_PAGES_PER_REQUEST})")
        logger.warning(f"‚ö†Ô∏è [STAGE 2] –†–∞–∑–±–∏–≤–∞–µ–º –Ω–∞ –±–∞—Ç—á–∏ –ø–æ {STAGE2_MAX_PAGES_PER_REQUEST} —Å—Ç—Ä–∞–Ω–∏—Ü...")

        # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –±–∞—Ç—á–∞–º–∏
        all_page_mappings = []
        for batch_start in range(0, len(doc_images_low), STAGE2_MAX_PAGES_PER_REQUEST):
            batch_end = min(batch_start + STAGE2_MAX_PAGES_PER_REQUEST, len(doc_images_low))
            batch_images = doc_images_low[batch_start:batch_end]
            batch_metadata = pages_metadata[batch_start:batch_end]
            batch_page_numbers = page_numbers[batch_start:batch_end] if page_numbers else list(range(batch_start + 1, batch_end + 1))

            logger.info(f"üìÑ [STAGE 2] –û–±—Ä–∞–±–æ—Ç–∫–∞ –±–∞—Ç—á–∞ —Å—Ç—Ä–∞–Ω–∏—Ü {batch_start+1}-{batch_end}...")

            # –ê–Ω–∞–ª–∏–∑–∏—Ä—É–µ–º –±–∞—Ç—á
            batch_mapping = await _analyze_relevance_batch(batch_metadata, batch_images, requirements, 0, batch_page_numbers)
            all_page_mappings.append(batch_mapping)

        # –û–±—ä–µ–¥–∏–Ω—è–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã –≤—Å–µ—Ö –±–∞—Ç—á–µ–π
        combined_mapping = {}
        for req in requirements:
            req_num = req['number']
            combined_pages = []
            for batch_mapping in all_page_mappings:
                if req_num in batch_mapping:
                    combined_pages.extend(batch_mapping[req_num])
            # –£–±–∏—Ä–∞–µ–º –¥—É–±–ª–∏–∫–∞—Ç—ã –∏ —Å–æ—Ä—Ç–∏—Ä—É–µ–º
            combined_mapping[req_num] = sorted(list(set(combined_pages)))
            logger.info(f"üìÑ [STAGE 2] Req {req_num}: —Å—Ç—Ä–∞–Ω–∏—Ü—ã {combined_mapping[req_num][:5]}{'...' if len(combined_mapping[req_num]) > 5 else ''}")

        logger.info(f"‚úÖ [STAGE 2] –ü–æ—Å—Ç—Ä–æ–µ–Ω mapping –¥–ª—è {len(combined_mapping)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π")
        return combined_mapping

    # –ï—Å–ª–∏ —Å—Ç—Ä–∞–Ω–∏—Ü –Ω–µ–º–Ω–æ–≥–æ - –∞–Ω–∞–ª–∏–∑–∏—Ä—É–µ–º –æ–¥–Ω–∏–º –∑–∞–ø—Ä–æ—Å–æ–º
    return await _analyze_relevance_batch(pages_metadata, doc_images_low, requirements, 0, page_numbers or list(range(1, len(doc_images_low)+1)))


async def _analyze_relevance_batch(
    batch_metadata: List[Dict[str, Any]],
    batch_images: List[str],
    requirements: List[Dict[str, Any]],
    offset: int = 0,
    page_numbers: Optional[List[int]] = None
) -> Dict[int, List[int]]:
    """
    –í—Å–ø–æ–º–æ–≥–∞—Ç–µ–ª—å–Ω–∞—è —Ñ—É–Ω–∫—Ü–∏—è –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ –±–∞—Ç—á–∞ —Å—Ç—Ä–∞–Ω–∏—Ü.
    offset - —Å–º–µ—â–µ–Ω–∏–µ –Ω–æ–º–µ—Ä–æ–≤ —Å—Ç—Ä–∞–Ω–∏—Ü –¥–ª—è –∫–æ—Ä—Ä–µ–∫—Ç–Ω–æ–π –Ω—É–º–µ—Ä–∞—Ü–∏–∏
    """
    # –§–æ—Ä–º–∏—Ä—É–µ–º –æ–ø–∏—Å–∞–Ω–∏–µ —Å—Ç—Ä–∞–Ω–∏—Ü
    pages_description = "\n".join([
        f"–°—Ç—Ä–∞–Ω–∏—Ü–∞ {p['page']}: {p.get('title', 'N/A')} [{p.get('section', 'N/A')}] - {p.get('type', 'N/A')}"
        for p in batch_metadata
    ])

    # –§–æ—Ä–º–∏—Ä—É–µ–º —Å–ø–∏—Å–æ–∫ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π
    requirements_text = "\n".join([
        f"{req['number']}. [{req.get('section', '–û–±—â–∏–µ')}] {req['text'][:200]}..."
        for req in requirements
    ])

    # –ò—Å–ø–æ–ª—å–∑—É–µ–º –∑–∞–≥—Ä—É–∂–µ–Ω–Ω—ã–π –ø—Ä–æ–º–ø—Ç —Å –ø–æ–¥—Å—Ç–∞–Ω–æ–≤–∫–æ–π –ø–µ—Ä–µ–º–µ–Ω–Ω—ã—Ö
    prompt_text = STAGE_PROMPTS["stage2_relevance"].format(
        page_count=len(batch_images),
        pages_description=pages_description,
        requirements_text=requirements_text
    )

    content = [{
        "type": "text",
        "text": prompt_text
    }]

    # –î–æ–±–∞–≤–ª—è–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –≤ –í–´–°–û–ö–û–ú –∫–∞—á–µ—Å—Ç–≤–µ (gpt-5-mini –¥–µ—à–µ–≤–∞—è, –Ω–µ —ç–∫–æ–Ω–æ–º–∏–º)
    for idx, base64_image in enumerate(batch_images, 1):
        page_num = (page_numbers[idx - 1] if page_numbers and idx - 1 < len(page_numbers) else (offset + idx))
        content.append({
            "type": "text",
            "text": f"\n--- –°—Ç—Ä–∞–Ω–∏—Ü–∞ {page_num} ---"
        })
        content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/jpeg;base64,{base64_image}",
                "detail": STAGE2_DETAIL  # "high" –¥–ª—è —Ç–æ—á–Ω–æ—Å—Ç–∏
            }
        })

    try:
        response = await client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[{"role": "user", "content": content}],
            response_format={"type": "json_object"},
            max_completion_tokens=4000
        )

        data = json.loads(response.choices[0].message.content)
        page_mapping_list = data.get('page_mapping', [])

        # –ü—Ä–µ–æ–±—Ä–∞–∑—É–µ–º –≤ —Å–ª–æ–≤–∞—Ä—å {requirement_number: [pages]}
        page_mapping = {}
        for item in page_mapping_list:
            req_num = item.get('requirement_number')
            pages = item.get('relevant_pages', [])
            reason = item.get('reason', '')
            page_mapping[req_num] = pages
            logger.info(f"üìÑ [STAGE 2] Req {req_num}: —Å—Ç—Ä–∞–Ω–∏—Ü—ã {pages} ({reason})")

        logger.info(f"‚úÖ [STAGE 2] –ü–æ—Å—Ç—Ä–æ–µ–Ω mapping –¥–ª—è {len(page_mapping)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π")
        return page_mapping

    except Exception as e:
        logger.error(f"‚ùå [STAGE 2] –û—à–∏–±–∫–∞ –æ—Ü–µ–Ω–∫–∏ —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω–æ—Å—Ç–∏ –±–∞—Ç—á–∞: {e}")
        # Fallback: –≤—Å–µ —Å—Ç—Ä–∞–Ω–∏—Ü—ã –∏–∑ —ç—Ç–æ–≥–æ –±–∞—Ç—á–∞ –¥–ª—è –≤—Å–µ—Ö —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π
        logger.warning(f"‚ö†Ô∏è [STAGE 2] –ò—Å–ø–æ–ª—å–∑—É–µ–º fallback –¥–ª—è –±–∞—Ç—á–∞ - —Å—Ç—Ä–∞–Ω–∏—Ü—ã {offset+1}-{offset+len(batch_images)}")
        return {req['number']: list(range(offset + 1, offset + len(batch_images) + 1)) for req in requirements}


# ============================
# TEXT PREFILTER FOR STAGE 2
# ============================

def _extract_page_texts_quick(doc_content: bytes, max_pages: int = 200) -> List[str]:
    """–ë—ã—Å—Ç—Ä–æ –∏–∑–≤–ª–µ–∫–∞–µ—Ç —Ç–µ–∫—Å—Ç –ø–æ —Å—Ç—Ä–∞–Ω–∏—Ü–∞–º –±–µ–∑ OCR (–¥–ª—è –ø—Ä–µ—Ñ–∏–ª—å—Ç—Ä–∞)."""
    with PDFProcessor(doc_content, "temp.pdf") as processor:
        return processor.extract_text_pages(max_pages)


def _simple_candidate_pages(requirements: List[Dict[str, Any]], page_texts: List[str], per_req: int = 7, cap_total: int = 30) -> List[int]:
    """–ü—Ä–æ—Å—Ç–æ–π —Ç–µ–∫—Å—Ç–æ–≤—ã–π –ø—Ä–µ—Ñ–∏–ª—å—Ç—Ä: –≤—ã–±–∏—Ä–∞–µ—Ç top-k —Å—Ç—Ä–∞–Ω–∏—Ü –¥–ª—è –∫–∞–∂–¥–æ–≥–æ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –ø–æ —Å–æ–≤–ø–∞–¥–µ–Ω–∏—é –∫–ª—é—á–µ–≤—ã—Ö —Å–ª–æ–≤."""
    import re
    candidates: List[int] = []
    # –û—á–µ–Ω—å –ø—Ä–æ—Å—Ç–æ–π tokenizer
    def toks(s: str) -> List[str]:
        return re.findall(r"[A-Za-z–ê-–Ø–∞-—è–Å—ë0-9_-]{2,}", (s or "").lower())

    page_tokens = [toks(t) for t in page_texts]

    for req in requirements:
        words = toks(req.get('text', ''))
        if not words:
            continue
        scores: List[Tuple[int, int]] = []  # (page_index, score)
        word_set = set(words)
        for idx, p_tokens in enumerate(page_tokens):
            if not p_tokens:
                continue
            score = sum(1 for w in p_tokens if w in word_set)
            if score:
                scores.append((idx, score))
        scores.sort(key=lambda x: x[1], reverse=True)
        top = [i for (i, _) in scores[:per_req]]
        candidates.extend(top)

    # –£–Ω–∏–∫–∞–ª—å–Ω—ã–µ, –æ—Ç—Å–æ—Ä—Ç–∏—Ä–æ–≤–∞–Ω–Ω—ã–µ, –æ–≥—Ä–∞–Ω–∏—á–µ–Ω–Ω—ã–µ –ø–æ –∫–æ–ª–∏—á–µ—Å—Ç–≤—É
    uniq = sorted(list({i for i in candidates}))
    if not uniq:
        # fallback: –ø–µ—Ä–≤—ã–µ 20 —Å—Ç—Ä–∞–Ω–∏—Ü
        uniq = list(range(0, min(20, len(page_texts))))
    return [i + 1 for i in uniq[:cap_total]]  # 1-based


@retry(stop=stop_after_attempt(RETRY_MAX_ATTEMPTS), wait=wait_exponential(multiplier=RETRY_WAIT_EXPONENTIAL_MULTIPLIER, min=4, max=RETRY_WAIT_EXPONENTIAL_MAX))
async def find_contradictions(
    pages_metadata: List[Dict[str, Any]],
    doc_content: bytes,
    requirements: List[Dict[str, Any]],
    analyzed_reqs: List['RequirementAnalysis']
) -> str:
    """
    Stage 4: –ü–æ–∏—Å–∫ –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏–π –≤ –ø—Ä–æ–µ–∫—Ç–Ω–æ–π –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏.
    –ê–Ω–∞–ª–∏–∑–∏—Ä—É–µ—Ç –∫–ª—é—á–µ–≤—ã–µ —Å—Ç—Ä–∞–Ω–∏—Ü—ã –∏–∑ —Ä–∞–∑–Ω—ã—Ö —Ä–∞–∑–¥–µ–ª–æ–≤ –∏ –∏—â–µ—Ç –Ω–µ—Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤–∏—è.

    Returns: —Ç–µ–∫—Å—Ç–æ–≤—ã–π –æ—Ç—á–µ—Ç –æ –Ω–∞–π–¥–µ–Ω–Ω—ã—Ö –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏—è—Ö
    """
    logger.info(f"üîç [STAGE 4] –ù–∞—á–∞–ª–æ –ø–æ–∏—Å–∫–∞ –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏–π –≤ –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏...")

    # –ì—Ä—É–ø–ø–∏—Ä—É–µ–º —Å—Ç—Ä–∞–Ω–∏—Ü—ã –ø–æ —Ä–∞–∑–¥–µ–ª–∞–º
    sections = {}
    for page_meta in pages_metadata:
        section = page_meta.get('section', 'N/A')
        if section not in sections:
            sections[section] = []
        sections[section].append(page_meta)

    logger.info(f"üìä [STAGE 4] –ù–∞–π–¥–µ–Ω–æ —Ä–∞–∑–¥–µ–ª–æ–≤: {list(sections.keys())}")

    # –û—Ç–±–∏—Ä–∞–µ–º –∫–ª—é—á–µ–≤—ã–µ —Å—Ç—Ä–∞–Ω–∏—Ü—ã –∏–∑ –∫–∞–∂–¥–æ–≥–æ —Ä–∞–∑–¥–µ–ª–∞
    selected_pages = []
    for section, pages in sections.items():
        # –ë–µ—Ä–µ–º –ø–µ—Ä–≤—ã–µ N —Å—Ç—Ä–∞–Ω–∏—Ü –∏–∑ –∫–∞–∂–¥–æ–≥–æ —Ä–∞–∑–¥–µ–ª–∞
        sample = pages[:STAGE4_SAMPLE_PAGES_PER_SECTION]
        selected_pages.extend([p['page'] for p in sample])
        logger.info(f"üìÑ [STAGE 4] –†–∞–∑–¥–µ–ª {section}: –≤—ã–±—Ä–∞–Ω–æ {len(sample)} —Å—Ç—Ä–∞–Ω–∏—Ü")

    # –ò–∑–≤–ª–µ–∫–∞–µ–º –≤—ã–±—Ä–∞–Ω–Ω—ã–µ —Å—Ç—Ä–∞–Ω–∏—Ü—ã –≤ —Å—Ä–µ–¥–Ω–µ–º –∫–∞—á–µ—Å—Ç–≤–µ
    logger.info(f"üìÑ [STAGE 4] –ò–∑–≤–ª–µ—á–µ–Ω–∏–µ {len(selected_pages)} –∫–ª—é—á–µ–≤—ã—Ö —Å—Ç—Ä–∞–Ω–∏—Ü...")

    # –ò—Å–ø–æ–ª—å–∑—É–µ–º –Ω–æ–≤—ã–π PDFBatchProcessor –¥–ª—è –ø–∞—Ä–∞–ª–ª–µ–ª—å–Ω–æ–≥–æ –∏–∑–≤–ª–µ—á–µ–Ω–∏—è
    processor = PDFBatchProcessor(doc_content, "contradictions_analysis.pdf")
    image_bases = await processor.extract_pages_batch(selected_pages, STAGE4_DPI, STAGE4_QUALITY)

    doc_images = []
    for page_num, base64_image in zip(selected_pages, image_bases):
        if base64_image:  # –¢–æ–ª—å–∫–æ —É—Å–ø–µ—à–Ω—ã–µ –∏–∑–≤–ª–µ—á–µ–Ω–∏—è
            doc_images.append({'page': page_num, 'image': base64_image})

    logger.info(f"‚úÖ [STAGE 4] –ò–∑–≤–ª–µ—á–µ–Ω–æ {len(doc_images)} —Å—Ç—Ä–∞–Ω–∏—Ü")

    # –§–æ—Ä–º–∏—Ä—É–µ–º summary –ø—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–Ω—ã—Ö —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π
    requirements_summary = "\n".join([
        f"{r.number}. {r.requirement[:100]}... ‚Üí {r.status} (—É–≤–µ—Ä–µ–Ω–Ω–æ—Å—Ç—å: {r.confidence}%)"
        for r in analyzed_reqs[:20]  # –ü–µ—Ä–≤—ã–µ 20 –¥–ª—è —ç–∫–æ–Ω–æ–º–∏–∏ —Ç–æ–∫–µ–Ω–æ–≤
    ])

    # –ò—Å–ø–æ–ª—å–∑—É–µ–º –∑–∞–≥—Ä—É–∂–µ–Ω–Ω—ã–π –ø—Ä–æ–º–ø—Ç
    prompt_text = STAGE_PROMPTS["stage4_contradictions"].format(
        requirements_summary=requirements_summary,
        sections=', '.join(sections.keys()),
        page_count=len(doc_images)
    )

    # –§–æ—Ä–º–∏—Ä—É–µ–º –ø—Ä–æ–º–ø—Ç –¥–ª—è –ø–æ–∏—Å–∫–∞ –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏–π
    content = [{
        "type": "text",
        "text": prompt_text
    }]

    # –î–æ–±–∞–≤–ª—è–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
    for img_data in doc_images:
        content.append({
            "type": "text",
            "text": f"\n--- –°—Ç—Ä–∞–Ω–∏—Ü–∞ {img_data['page']} ---"
        })
        content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/jpeg;base64,{img_data['image']}",
                "detail": STAGE4_DETAIL
            }
        })

    try:
        response = await client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[{"role": "user", "content": content}],
            response_format={"type": "json_object"},
            max_completion_tokens=STAGE4_MAX_COMPLETION_TOKENS
        )

        result = json.loads(response.choices[0].message.content)
        contradictions = result.get('contradictions', [])
        summary = result.get('summary', '–ê–Ω–∞–ª–∏–∑ –∑–∞–≤–µ—Ä—à–µ–Ω')

        logger.info(f"‚úÖ [STAGE 4] –ù–∞–π–¥–µ–Ω–æ –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏–π: {len(contradictions)}")

        # –§–æ—Ä–º–∏—Ä—É–µ–º —Ç–µ–∫—Å—Ç–æ–≤—ã–π –æ—Ç—á–µ—Ç
        if not contradictions:
            return "‚úÖ –ü–†–û–¢–ò–í–û–†–ï–ß–ò–ô –ù–ï –û–ë–ù–ê–†–£–ñ–ï–ù–û\n\n–ü—Ä–æ–µ–∫—Ç–Ω–∞—è –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏—è –Ω–µ —Å–æ–¥–µ—Ä–∂–∏—Ç —è–≤–Ω—ã—Ö –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏–π –º–µ–∂–¥—É —Ä–∞–∑–¥–µ–ª–∞–º–∏."

        report = f"üîç –û–¢–ß–ï–¢ –û –ü–†–û–¢–ò–í–û–†–ï–ß–ò–Ø–•\n\n{summary}\n\n"
        report += "=" * 80 + "\n\n"

        for idx, contr in enumerate(contradictions, 1):
            severity_emoji = {"–∫—Ä–∏—Ç–∏—á–Ω–æ": "üî¥", "—Å—Ä–µ–¥–Ω–µ": "üü°", "–Ω–∏–∑–∫–æ": "üü¢"}.get(contr.get('severity', '—Å—Ä–µ–¥–Ω–µ'), "‚ö™")
            report += f"{idx}. {severity_emoji} {contr.get('type', '–ù–µ—Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤–∏–µ').upper()}\n"
            report += f"   –ö—Ä–∏—Ç–∏—á–Ω–æ—Å—Ç—å: {contr.get('severity', '—Å—Ä–µ–¥–Ω–µ')}\n"
            report += f"   –û–ø–∏—Å–∞–Ω–∏–µ: {contr.get('description', 'N/A')}\n"
            report += f"   –°—Ç—Ä–∞–Ω–∏—Ü—ã: {', '.join(map(str, contr.get('pages', [])))}\n"
            report += f"   –†–µ–∫–æ–º–µ–Ω–¥–∞—Ü–∏—è: {contr.get('recommendation', '–¢—Ä–µ–±—É–µ—Ç—Å—è —É—Ç–æ—á–Ω–µ–Ω–∏–µ')}\n\n"

        return report

    except Exception as e:
        logger.error(f"‚ùå [STAGE 4] –û—à–∏–±–∫–∞ –ø–æ–∏—Å–∫–∞ –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏–π: {e}")
        return f"‚ö†Ô∏è –û–®–ò–ë–ö–ê –ê–ù–ê–õ–ò–ó–ê –ü–†–û–¢–ò–í–û–†–ï–ß–ò–ô\n\n–ù–µ —É–¥–∞–ª–æ—Å—å –≤—ã–ø–æ–ª–Ω–∏—Ç—å –∞–Ω–∞–ª–∏–∑: {str(e)}"


def normalize_status_confidence(analysis: Dict[str, Any]) -> Dict[str, Any]:
    """
    –ù–æ—Ä–º–∞–ª–∏–∑—É–µ—Ç –Ω–µ—Å–æ–≥–ª–∞—Å–æ–≤–∞–Ω–Ω–æ—Å—Ç–∏ –º–µ–∂–¥—É status –∏ confidence.
    
    –ü—Ä–∞–≤–∏–ª–∞:
    - "–ü–æ–ª–Ω–æ—Å—Ç—å—é –∏—Å–ø–æ–ª–Ω–µ–Ω–æ" ‚Üí confidence >= 70%
    - "–ß–∞—Å—Ç–∏—á–Ω–æ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ" ‚Üí confidence 40-80%
    - "–ù–µ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ" ‚Üí confidence >= 60% (—É–≤–µ—Ä–µ–Ω—ã, —á—Ç–æ –ù–ï –∏—Å–ø–æ–ª–Ω–µ–Ω–æ)
    - "–¢—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è" ‚Üí confidence <= 40%
    """
    status = analysis.get('status', '')
    confidence = analysis.get('confidence', 0)
    
    # –ù–æ—Ä–º–∞–ª–∏–∑—É–µ–º confidence –≤ –¥–∏–∞–ø–∞–∑–æ–Ω 0-100
    if confidence < 0:
        confidence = 0
    elif confidence > 100:
        confidence = 100
    
    # –ü—Ä–æ–≤–µ—Ä—è–µ–º –∏ –∫–æ—Ä—Ä–µ–∫—Ç–∏—Ä—É–µ–º –Ω–µ—Å–æ–≥–ª–∞—Å–æ–≤–∞–Ω–Ω–æ—Å—Ç–∏
    if status == "–ü–æ–ª–Ω–æ—Å—Ç—å—é –∏—Å–ø–æ–ª–Ω–µ–Ω–æ":
        if confidence < 70:
            # –ù–∏–∑–∫–∞—è —É–≤–µ—Ä–µ–Ω–Ω–æ—Å—Ç—å –¥–ª—è "–ø–æ–ª–Ω–æ—Å—Ç—å—é –∏—Å–ø–æ–ª–Ω–µ–Ω–æ" - –º–µ–Ω—è–µ–º —Å—Ç–∞—Ç—É—Å
            logger.warning(f"‚ö†Ô∏è –ù–µ—Å–æ–≥–ª–∞—Å–æ–≤–∞–Ω–Ω–æ—Å—Ç—å: status='–ü–æ–ª–Ω–æ—Å—Ç—å—é –∏—Å–ø–æ–ª–Ω–µ–Ω–æ', confidence={confidence}. –ò—Å–ø—Ä–∞–≤–ª—è–µ–º —Å—Ç–∞—Ç—É—Å –Ω–∞ '–ß–∞—Å—Ç–∏—á–Ω–æ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ'")
            analysis['status'] = "–ß–∞—Å—Ç–∏—á–Ω–æ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ"
            
    elif status == "–¢—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è":
        if confidence > 40:
            # –í—ã—Å–æ–∫–∞—è —É–≤–µ—Ä–µ–Ω–Ω–æ—Å—Ç—å –¥–ª—è "—Ç—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è" - –∫–æ—Ä—Ä–µ–∫—Ç–∏—Ä—É–µ–º
            logger.warning(f"‚ö†Ô∏è –ù–µ—Å–æ–≥–ª–∞—Å–æ–≤–∞–Ω–Ω–æ—Å—Ç—å: status='–¢—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è', confidence={confidence}. –°–Ω–∏–∂–∞–µ–º confidence –¥–æ 30%")
            analysis['confidence'] = 30
            
    elif status == "–ß–∞—Å—Ç–∏—á–Ω–æ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ":
        # –ü—Ä–∏–≤–æ–¥–∏–º –∫ —Ä–∞–∑—É–º–Ω—ã–º –≥—Ä–∞–Ω–∏—Ü–∞–º
        if confidence < 40:
            analysis['confidence'] = 40
        elif confidence > 80:
            # –°–ª–∏—à–∫–æ–º –≤—ã—Å–æ–∫–∞—è —É–≤–µ—Ä–µ–Ω–Ω–æ—Å—Ç—å - –≤–æ–∑–º–æ–∂–Ω–æ "–ø–æ–ª–Ω–æ—Å—Ç—å—é –∏—Å–ø–æ–ª–Ω–µ–Ω–æ"
            logger.warning(f"‚ö†Ô∏è –ù–µ—Å–æ–≥–ª–∞—Å–æ–≤–∞–Ω–Ω–æ—Å—Ç—å: status='–ß–∞—Å—Ç–∏—á–Ω–æ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ', confidence={confidence}. –ú–µ–Ω—è–µ–º –Ω–∞ '–ü–æ–ª–Ω–æ—Å—Ç—å—é –∏—Å–ø–æ–ª–Ω–µ–Ω–æ'")
            analysis['status'] = "–ü–æ–ª–Ω–æ—Å—Ç—å—é –∏—Å–ø–æ–ª–Ω–µ–Ω–æ"
            
    elif status == "–ù–µ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ":
        # –î–ª—è "–Ω–µ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ" confidence –ø–æ–∫–∞–∑—ã–≤–∞–µ—Ç —É–≤–µ—Ä–µ–Ω–Ω–æ—Å—Ç—å –≤ –æ—Ç—Å—É—Ç—Å—Ç–≤–∏–∏
        if confidence < 50:
            # –ù–∏–∑–∫–∞—è —É–≤–µ—Ä–µ–Ω–Ω–æ—Å—Ç—å - –≤–æ–∑–º–æ–∂–Ω–æ —Ç—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è
            logger.warning(f"‚ö†Ô∏è –ù–µ—Å–æ–≥–ª–∞—Å–æ–≤–∞–Ω–Ω–æ—Å—Ç—å: status='–ù–µ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ', confidence={confidence}. –ú–µ–Ω—è–µ–º –Ω–∞ '–¢—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è'")
            analysis['status'] = "–¢—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è"
            analysis['confidence'] = confidence
    
    return analysis


def get_analysis_system_prompt(stage: str, req_type: str) -> str:
    """
    –í–æ–∑–≤—Ä–∞—â–∞–µ—Ç system prompt –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏.
    """
    stage_prompt = PROMPTS.get(stage, PROMPTS["–§–≠"])
    # –í system_prompt_template –µ—Å—Ç—å –Ω–µ–æ–±—Ä–∞–º–ª–µ–Ω–Ω—ã–µ —Ñ–∏–≥—É—Ä–Ω—ã–µ —Å–∫–æ–±–∫–∏ (JSON-–ø—Ä–∏–º–µ—Ä—ã),
    # –ø–æ—ç—Ç–æ–º—É –∏—Å–ø–æ–ª—å–∑—É–µ–º –±–µ–∑–æ–ø–∞—Å–Ω—É—é –∑–∞–º–µ–Ω—É —Ç–æ–ª—å–∫–æ –ø–ª–µ–π—Å—Ö–æ–ª–¥–µ—Ä–∞ {req_type}.
    system_template_raw = STAGE_PROMPTS["system_prompt_template"]
    system_template = system_template_raw.replace("{req_type}", req_type)

    return f"{stage_prompt}\n\n{system_template}"


async def analyze_batch_with_high_detail(
    system_prompt: str,
    doc_content: bytes,
    page_numbers: List[int],
    requirements_batch: List[Dict[str, Any]],
    request: Request,
    pages_metadata: Optional[List[Dict[str, Any]]] = None
) -> List['RequirementAnalysis']:
    """
    Stage 3: –î–µ—Ç–∞–ª—å–Ω—ã–π –∞–Ω–∞–ª–∏–∑ –ø–∞–∫–µ—Ç–∞ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π —Å –í–´–°–û–ö–ò–ú —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–µ–º —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω—ã—Ö —Å—Ç—Ä–∞–Ω–∏—Ü.
    
    Args:
        pages_metadata: –ú–µ—Ç–∞–¥–∞–Ω–Ω—ã–µ —Å—Ç—Ä–∞–Ω–∏—Ü —Å –Ω–æ–º–µ—Ä–∞–º–∏ –ª–∏—Å—Ç–æ–≤ –∏–∑ Stage 1
    """
    if await request.is_disconnected():
        logger.warning(f"‚ö†Ô∏è [STAGE 3] Client disconnected")
        return []

    batch_ids = [req['trace_id'] for req in requirements_batch]
    logger.info(f"üîç [STAGE 3] –ê–Ω–∞–ª–∏–∑ {len(requirements_batch)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π –Ω–∞ —Å—Ç—Ä–∞–Ω–∏—Ü–∞—Ö {page_numbers[:5]}{'...' if len(page_numbers) > 5 else ''}")

    # –ü—Ä–æ–≤–µ—Ä–∫–∞ –ª–∏–º–∏—Ç–∞ —Å—Ç—Ä–∞–Ω–∏—Ü (–∑–∞—â–∏—Ç–∞ –æ—Ç 429 rate limit)
    if len(page_numbers) > STAGE3_MAX_PAGES_PER_REQUEST:
        logger.warning(f"‚ö†Ô∏è [STAGE 3] –°–ª–∏—à–∫–æ–º –º–Ω–æ–≥–æ —Å—Ç—Ä–∞–Ω–∏—Ü ({len(page_numbers)} > {STAGE3_MAX_PAGES_PER_REQUEST})")
        logger.warning(f"‚ö†Ô∏è [STAGE 3] –†–∞–∑–±–∏–≤–∞–µ–º —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –Ω–∞ –ø–æ–¥–≥—Ä—É–ø–ø—ã –ø–æ —Å—Ç—Ä–∞–Ω–∏—Ü–∞–º...")

        # –†–∞–∑–±–∏–≤–∞–µ–º page_numbers –Ω–∞ —á–∞–Ω–∫–∏
        all_results = []
        for i in range(0, len(page_numbers), STAGE3_MAX_PAGES_PER_REQUEST):
            chunk_pages = page_numbers[i:i + STAGE3_MAX_PAGES_PER_REQUEST]
            logger.info(f"üìÑ [STAGE 3] –û–±—Ä–∞–±–æ—Ç–∫–∞ –ø–æ–¥–≥—Ä—É–ø–ø—ã —Å—Ç—Ä–∞–Ω–∏—Ü {i+1}-{min(i+STAGE3_MAX_PAGES_PER_REQUEST, len(page_numbers))}")

            # –†–µ–∫—É—Ä—Å–∏–≤–Ω—ã–π –≤—ã–∑–æ–≤ —Å –º–µ–Ω—å—à–∏–º –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ–º —Å—Ç—Ä–∞–Ω–∏—Ü
            chunk_results = await analyze_batch_with_high_detail(
                system_prompt=system_prompt,
                doc_content=doc_content,
                page_numbers=chunk_pages,
                requirements_batch=requirements_batch,
                request=request,
                pages_metadata=pages_metadata
            )

            if not chunk_results:
                return []  # Client disconnected

            all_results.extend(chunk_results)

        # –£–±–∏—Ä–∞–µ–º –¥—É–±–ª–∏–∫–∞—Ç—ã –ø–æ –Ω–æ–º–µ—Ä—É —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è
        seen = set()
        unique_results = []
        for result in all_results:
            if result.number not in seen:
                seen.add(result.number)
                unique_results.append(result)

        return unique_results

    # –ò–∑–≤–ª–µ–∫–∞–µ–º —Ç–æ–ª—å–∫–æ —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω—ã–µ —Å—Ç—Ä–∞–Ω–∏—Ü—ã –≤ –≤—ã—Å–æ–∫–æ–º –∫–∞—á–µ—Å—Ç–≤–µ
    logger.info(f"üìÑ [STAGE 3] –ò–∑–≤–ª–µ—á–µ–Ω–∏–µ {len(page_numbers)} —Å—Ç—Ä–∞–Ω–∏—Ü –≤ –≤—ã—Å–æ–∫–æ–º –∫–∞—á–µ—Å—Ç–≤–µ...")

    # –ò—Å–ø–æ–ª—å–∑—É–µ–º PDFBatchProcessor –¥–ª—è –ø–∞—Ä–∞–ª–ª–µ–ª—å–Ω–æ–≥–æ –∏–∑–≤–ª–µ—á–µ–Ω–∏—è –≤ –≤—ã—Å–æ–∫–æ–º –∫–∞—á–µ—Å—Ç–≤–µ
    processor = PDFBatchProcessor(doc_content, "stage3_analysis.pdf")
    doc_images_high = await processor.extract_pages_batch(
        page_numbers, STAGE3_DPI, STAGE3_QUALITY, max_concurrent=3  # –ú–µ–Ω—å—à–µ –æ–¥–Ω–æ–≤—Ä–µ–º–µ–Ω–Ω—ã—Ö –¥–ª—è –≤—ã—Å–æ–∫–æ–≥–æ –∫–∞—á–µ—Å—Ç–≤–∞
    )

    logger.info(f"‚úÖ [STAGE 3] –ò–∑–≤–ª–µ—á–µ–Ω–æ {len([img for img in doc_images_high if img])} —Å—Ç—Ä–∞–Ω–∏—Ü –≤ –≤—ã—Å–æ–∫–æ–º –∫–∞—á–µ—Å—Ç–≤–µ")

    # –§–æ—Ä–º–∏—Ä—É–µ–º —Å–ø–∏—Å–æ–∫ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π
    requirements_text = "\n\n".join([
        f"–¢—Ä–µ–±–æ–≤–∞–Ω–∏–µ {req['number']} [{req.get('section', '–û–±—â–∏–µ')}]:\n{req['text']}"
        for req in requirements_batch
    ])

    # –°–æ–∑–¥–∞–µ–º mapping PDF —Å—Ç—Ä–∞–Ω–∏—Ü–∞ ‚Üí –ù–æ–º–µ—Ä –ª–∏—Å—Ç–∞ –ø—Ä–æ–µ–∫—Ç–∞
    page_to_sheet_mapping = {}
    if pages_metadata:
        for page_meta in pages_metadata:
            pdf_page = page_meta.get('page')
            sheet_num = page_meta.get('sheet_number', f"–°—Ç—Ä.{pdf_page}")
            if pdf_page:
                page_to_sheet_mapping[pdf_page] = sheet_num
    
    # –§–æ—Ä–º–∏—Ä—É–µ–º —Å–ø–∏—Å–æ–∫ –¥–æ—Å—Ç—É–ø–Ω—ã—Ö —Å—Ç—Ä–∞–Ω–∏—Ü —Å –Ω–æ–º–µ—Ä–∞–º–∏ –ª–∏—Å—Ç–æ–≤
    if page_to_sheet_mapping:
        available_pages_list = []
        for page_num in page_numbers:
            sheet_num = page_to_sheet_mapping.get(page_num, f"–°—Ç—Ä.{page_num}")
            available_pages_list.append(f"PDF —Å—Ç—Ä.{page_num} (–õ–∏—Å—Ç {sheet_num})")
        available_pages_str = ", ".join(available_pages_list)
        logger.info(f"üìÑ [STAGE 3] –î–æ—Å—Ç—É–ø–Ω—ã–µ –ª–∏—Å—Ç—ã: {available_pages_str[:200]}...")
    else:
        # Fallback –µ—Å–ª–∏ –Ω–µ—Ç –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö
        available_pages_str = ", ".join(map(str, page_numbers))
        logger.warning(f"‚ö†Ô∏è [STAGE 3] –ù–µ—Ç –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö –æ –ª–∏—Å—Ç–∞—Ö, –∏—Å–ø–æ–ª—å–∑—É–µ–º PDF —Å—Ç—Ä–∞–Ω–∏—Ü—ã")

    # –ò—Å–ø–æ–ª—å–∑—É–µ–º –∑–∞–≥—Ä—É–∂–µ–Ω–Ω—ã–π –ø—Ä–æ–º–ø—Ç
    prompt_text = STAGE_PROMPTS["stage3_analysis"].format(
        requirements_text=requirements_text,
        requirements_count=len(requirements_batch),
        available_pages=available_pages_str,
        page_count=len(page_numbers)
    )

    content = [{
        "type": "text",
        "text": prompt_text
    }]

    # –î–æ–±–∞–≤–ª—è–µ–º –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è –≤ –≤—ã—Å–æ–∫–æ–º –∫–∞—á–µ—Å—Ç–≤–µ —Å –Ω–æ–º–µ—Ä–∞–º–∏ –ª–∏—Å—Ç–æ–≤
    for idx, base64_image in enumerate(doc_images_high, 1):
        page_num = page_numbers[idx - 1] if idx <= len(page_numbers) else idx
        
        # –ü–æ–ª—É—á–∞–µ–º –Ω–æ–º–µ—Ä –ª–∏—Å—Ç–∞ –ø—Ä–æ–µ–∫—Ç–∞ –¥–ª—è –ø–æ–¥–ø–∏—Å–∏
        if page_to_sheet_mapping and page_num in page_to_sheet_mapping:
            sheet_num = page_to_sheet_mapping[page_num]
            page_label = f"PDF —Å—Ç—Ä.{page_num} (–õ–∏—Å—Ç {sheet_num})"
        else:
            page_label = f"–°—Ç—Ä–∞–Ω–∏—Ü–∞ {page_num}"
        
        content.append({
            "type": "text",
            "text": f"\n--- {page_label} ---"
        })
        content.append({
            "type": "image_url",
            "image_url": {
                "url": f"data:image/jpeg;base64,{base64_image}",
                "detail": STAGE3_DETAIL  # –í—ã—Å–æ–∫–æ–µ –∫–∞—á–µ—Å—Ç–≤–æ
            }
        })

    try:
        response = await client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": content}
            ],
            response_format={"type": "json_object"},  # –ü—Ä–∏–Ω—É–¥–∏—Ç–µ–ª—å–Ω—ã–π JSON
            max_completion_tokens=STAGE3_MAX_COMPLETION_TOKENS
        )

        response_text = response.choices[0].message.content
        refusal = response.choices[0].message.refusal

        # –ü—Ä–æ–≤–µ—Ä–∫–∞ –Ω–∞ None –∏–ª–∏ refusal
        if response_text is None or refusal:
            logger.error(f"‚ùå [STAGE 3] Model refused to respond!")
            logger.error(f"Refusal message: {refusal}")
            logger.error(f"Finish reason: {response.choices[0].finish_reason}")
            logger.error(f"Requirements in batch: {[req['number'] for req in requirements_batch]}")
            logger.error(f"Requirements text preview: {[req['text'][:100] for req in requirements_batch]}")

            # –ï—Å–ª–∏ –±–∞—Ç—á –±–æ–ª—å—à–µ 1 –∏ –≤–∫–ª—é—á–µ–Ω retry - –ø—Ä–æ–±—É–µ–º –ø–æ –æ–¥–Ω–æ–º—É
            if STAGE3_RETRY_ON_REFUSAL and len(requirements_batch) > 1:
                logger.warning(f"‚ö†Ô∏è [STAGE 3] Retry: analyzing {len(requirements_batch)} requirements individually...")
                all_results = []
                for single_req in requirements_batch:
                    try:
                        single_result = await analyze_batch_with_high_detail(
                            system_prompt=system_prompt,
                            doc_content=doc_content,
                            page_numbers=page_numbers,
                            requirements_batch=[single_req],
                            request=request,
                            pages_metadata=pages_metadata
                        )
                        all_results.extend(single_result)
                    except Exception as e:
                        logger.error(f"‚ùå [STAGE 3] Failed to analyze requirement {single_req['number']}: {e}")
                        all_results.append(RequirementAnalysis(
                            number=single_req['number'],
                            requirement=single_req['text'],
                            status="–¢—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è",
                            confidence=0,
                            solution_description="–û—à–∏–±–∫–∞ –ø—Ä–∏ –∏–Ω–¥–∏–≤–∏–¥—É–∞–ª—å–Ω–æ–º –∞–Ω–∞–ª–∏–∑–µ",
                            reference="-",
                            discrepancies=str(e),
                            section=single_req.get('section')
                        ))
                return all_results

            # –í–æ–∑–≤—Ä–∞—â–∞–µ–º –∑–∞–≥–ª—É—à–∫–∏ –¥–ª—è –≤—Å–µ—Ö —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π
            return [
                RequirementAnalysis(
                    number=req['number'],
                    requirement=req['text'],
                    status="–¢—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è",
                    confidence=0,
                    solution_description="–ú–æ–¥–µ–ª—å –æ—Ç–∫–∞–∑–∞–ª–∞—Å—å –∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞—Ç—å",
                    reference="-",
                    discrepancies=f"Content filter: {refusal or 'Response is None'}",
                    section=req.get('section')
                )
                for req in requirements_batch
            ]

        logger.info(f"üìÑ [STAGE 3] Response preview: {response_text[:LOG_RESPONSE_PREVIEW_LENGTH]}...")

        # –ü–∞—Ä—Å–∏–º JSON
        json_start = response_text.find('{')
        json_end = response_text.rfind('}') + 1
        if json_start != -1 and json_end > json_start:
            json_str = response_text[json_start:json_end]
            data = json.loads(json_str)

            # –ü–æ–¥–¥–µ—Ä–∂–∫–∞ –¥–≤—É—Ö —Ñ–æ—Ä–º–∞—Ç–æ–≤: { analyses: [...] } –ò–õ–ò –æ–¥–∏–Ω–æ—á–Ω—ã–π –æ–±—ä–µ–∫—Ç –∞–Ω–∞–ª–∏–∑–∞
            if 'analyses' in data:
                analyses = data.get('analyses', [])
            else:
                # –ü—ã—Ç–∞–µ–º—Å—è –∏–Ω—Ç–µ—Ä–ø—Ä–µ—Ç–∏—Ä–æ–≤–∞—Ç—å –∫–∞–∫ –µ–¥–∏–Ω–∏—á–Ω—ã–π –æ–±—ä–µ–∫—Ç –∞–Ω–∞–ª–∏–∑–∞
                analyses = [data]

            # –ù–æ—Ä–º–∞–ª–∏–∑—É–µ–º —Ç–∏–ø: –µ—Å–ª–∏ analyses ‚Äî —Å–ª–æ–≤–∞—Ä—å, –æ–±–æ—Ä–∞—á–∏–≤–∞–µ–º –µ–≥–æ –≤ —Å–ø–∏—Å–æ–∫
            if isinstance(analyses, dict):
                analyses = [analyses]
            req_map = {req['number']: req for req in requirements_batch}

            results = []
            for analysis in analyses:
                req_num = analysis.get('number')
                if req_num in req_map:
                    req = req_map[req_num]
                    # –ù–æ—Ä–º–∞–ª–∏–∑—É–µ–º –Ω–µ—Å–æ–≥–ª–∞—Å–æ–≤–∞–Ω–Ω–æ—Å—Ç–∏ –º–µ–∂–¥—É status –∏ confidence
                    normalized_analysis = normalize_status_confidence(analysis)

                    # üîß –ö–†–ò–¢–ò–ß–ï–°–ö–ò –í–ê–ñ–ù–û: –í–∞–ª–∏–¥–∞—Ü–∏—è –∏ –Ω–æ—Ä–º–∞–ª–∏–∑–∞—Ü–∏—è reference –∫ –¢–û–õ–¨–ö–û –¶–ò–§–†–û–í–û–ú–£ —Ñ–æ—Ä–º–∞—Ç—É
                    reference = normalized_analysis.get('reference', '-')
                    if reference and reference != "-":
                        # –ù–æ—Ä–º–∞–ª–∏–∑—É–µ–º reference –∫ —Ü–∏—Ñ—Ä–æ–≤–æ–º—É —Ñ–æ—Ä–º–∞—Ç—É –µ—Å–ª–∏ —Å–æ–¥–µ—Ä–∂–∏—Ç –±—É–∫–≤—É
                        import re
                        if re.search(r'[–∞-—è–ê-–Ø—ë–Åa-zA-Z]', reference):
                            logger.warning(f"‚ö†Ô∏è [STAGE 3] Req {req_num}: reference —Å–æ–¥–µ—Ä–∂–∏—Ç –±—É–∫–≤—ã: '{reference}'")
                            normalized_ref = normalize_sheet_number_to_digit(reference)
                            normalized_analysis['reference'] = normalized_ref
                            logger.info(f"üìã [STAGE 3] Req {req_num}: –Ω–æ—Ä–º–∞–ª–∏–∑–æ–≤–∞–Ω–æ reference '{reference}' ‚Üí '{normalized_ref}'")

                    results.append(RequirementAnalysis(
                        **normalized_analysis,
                        section=req.get('section')
                    ))

            logger.info(f"‚úÖ [STAGE 3] –ü—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–æ {len(results)}/{len(requirements_batch)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π")

            # –î–æ–±–∞–≤–ª—è–µ–º –∑–∞–≥–ª—É—à–∫–∏ –¥–ª—è –ø—Ä–æ–ø—É—â–µ–Ω–Ω—ã—Ö
            if len(results) < len(requirements_batch):
                analyzed_numbers = {r.number for r in results}
                missing = [req['number'] for req in requirements_batch if req['number'] not in analyzed_numbers]
                logger.warning(f"‚ö†Ô∏è [STAGE 3] –ü—Ä–æ–ø—É—â–µ–Ω—ã —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è: {missing}")
                for req in requirements_batch:
                    if req['number'] not in analyzed_numbers:
                        results.append(RequirementAnalysis(
                            number=req['number'],
                            requirement=req['text'],
                            status="–¢—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è",
                            confidence=0,
                            solution_description="–ù–µ –ø—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–æ",
                            reference="-",
                            discrepancies="–û—Ç—Å—É—Ç—Å—Ç–≤—É–µ—Ç –≤ –æ—Ç–≤–µ—Ç–µ –º–æ–¥–µ–ª–∏",
                            section=req.get('section')
                        ))

            results.sort(key=lambda r: r.number)
            return results
        else:
            logger.error(f"‚ùå [STAGE 3] No JSON in response. Full response: {response_text[:500]}")
            raise ValueError("No JSON found in response")

    except RateLimitError as e:
        error_msg = str(e)
        logger.error(f"‚ùå [STAGE 3] Rate limit exceeded: {error_msg}")

        # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —ç—Ç–æ –ø—Ä–µ–≤—ã—à–µ–Ω–∏–µ TPM –∏–∑-–∑–∞ –±–æ–ª—å—à–æ–≥–æ –∑–∞–ø—Ä–æ—Å–∞
        if "tokens per min" in error_msg.lower() and len(page_numbers) > 10:
            logger.warning(f"‚ö†Ô∏è [STAGE 3] –ü—Ä–µ–≤—ã—à–µ–Ω –ª–∏–º–∏—Ç —Ç–æ–∫–µ–Ω–æ–≤ –∏–∑-–∑–∞ {len(page_numbers)} —Å—Ç—Ä–∞–Ω–∏—Ü")
            logger.warning(f"‚ö†Ô∏è [STAGE 3] –ê–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ —Ä–∞–∑–±–∏–≤–∞–µ–º –Ω–∞ –º–µ–Ω—å—à–∏–µ —á–∞—Å—Ç–∏...")

            # –†–∞–∑–±–∏–≤–∞–µ–º —Å—Ç—Ä–∞–Ω–∏—Ü—ã –ø–æ–ø–æ–ª–∞–º
            mid = len(page_numbers) // 2
            chunk1_pages = page_numbers[:mid]
            chunk2_pages = page_numbers[mid:]

            all_results = []
            for chunk_pages in [chunk1_pages, chunk2_pages]:
                chunk_results = await analyze_batch_with_high_detail(
                    system_prompt=system_prompt,
                    doc_content=doc_content,
                    page_numbers=chunk_pages,
                    requirements_batch=requirements_batch,
                    request=request,
                    pages_metadata=pages_metadata
                )
                if not chunk_results:
                    return []
                all_results.extend(chunk_results)

            # –£–±–∏—Ä–∞–µ–º –¥—É–±–ª–∏–∫–∞—Ç—ã
            seen = set()
            unique_results = []
            for result in all_results:
                if result.number not in seen:
                    seen.add(result.number)
                    unique_results.append(result)

            return unique_results

        # –û–±—ã—á–Ω–∞—è 429 –æ—à–∏–±–∫–∞ - retry –¥–µ–∫–æ—Ä–∞—Ç–æ—Ä –æ–±—Ä–∞–±–æ—Ç–∞–µ—Ç
        raise

    except Exception as e:
        logger.error(f"‚ùå [STAGE 3] –û—à–∏–±–∫–∞ –∞–Ω–∞–ª–∏–∑–∞: {e}")
        if 'response_text' in locals() and response_text is not None:
            logger.error(f"Full response text: {response_text[:1000]}")
        else:
            logger.error("Response text is None or not available")
        return [
            RequirementAnalysis(
                number=req['number'],
                requirement=req['text'],
                status="–¢—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è",
                confidence=0,
                solution_description="–û—à–∏–±–∫–∞ –∞–Ω–∞–ª–∏–∑–∞",
                reference="-",
                discrepancies=str(e),
                section=req.get('section')
            )
            for req in requirements_batch
        ]




# ============================
# –í–°–ü–û–ú–û–ì–ê–¢–ï–õ–¨–ù–´–ï –§–£–ù–ö–¶–ò–ò –î–õ–Ø –¢–ó/–¢–£
# ============================

@retry(stop=stop_after_attempt(RETRY_MAX_ATTEMPTS), wait=wait_exponential(multiplier=RETRY_WAIT_EXPONENTIAL_MULTIPLIER, min=4, max=RETRY_WAIT_EXPONENTIAL_MAX))
async def extract_text_from_pdf(content: bytes, filename: str) -> str:
    """–ò–∑–≤–ª–µ–∫–∞–µ—Ç —Ç–µ–∫—Å—Ç –∏–∑ PDF. –ò—Å–ø–æ–ª—å–∑—É–µ—Ç OCR –µ—Å–ª–∏ –Ω–µ—Ç —Ç–µ–∫—Å—Ç–æ–≤–æ–≥–æ —Å–ª–æ—è."""
    import base64
    from PIL import Image
    import io

    text = ""

    with PDFProcessor(content, filename) as processor:
        # Mixed-mode: –ø–æ—Å—Ç—Ä–∞–Ω–∏—á–Ω–æ
        for page_index in range(processor.page_count):
            page = processor.get_page(page_index + 1)  # 1-based
            page_text = page.get_text() or ""

            if page_text.strip():
                text += page_text + "\n\n"
                continue

            # OCR —Ç–æ–ª—å–∫–æ –¥–ª—è –ø—É—Å—Ç—ã—Ö –ø–æ —Ç–µ–∫—Å—Ç—É —Å—Ç—Ä–∞–Ω–∏—Ü
            logger.info(f"üìÑ OCR —Å—Ç—Ä–∞–Ω–∏—Ü—ã {page_index + 1}/{processor.page_count} –∏–∑ {filename}")

            # –ò—Å–ø–æ–ª—å–∑—É–µ–º –Ω–æ–≤—ã–π –ø—Ä–æ—Ü–µ—Å—Å–æ—Ä –¥–ª—è –∏–∑–≤–ª–µ—á–µ–Ω–∏—è –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è
            images = processor.extract_pages_as_images([page_index + 1], 100, 70)
            if not images:
                continue

            base64_image = images[0]

            response = await client.chat.completions.create(
                model=OPENAI_MODEL,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": "–ò–∑–≤–ª–µ–∫–∏ –≤–µ—Å—å —Ç–µ–∫—Å—Ç —Å —ç—Ç–æ–≥–æ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è. –°–æ—Ö—Ä–∞–Ω–∏ —Å—Ç—Ä—É–∫—Ç—É—Ä—É, –Ω–æ–º–µ—Ä–∞ –ø—É–Ω–∫—Ç–æ–≤, —Ç–∞–±–ª–∏—Ü—ã. –í–µ—Ä–Ω–∏ —Ç–æ–ª—å–∫–æ —Ç–µ–∫—Å—Ç –±–µ–∑ –∫–æ–º–º–µ–Ω—Ç–∞—Ä–∏–µ–≤."
                            },
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}
                            }
                        ]
                    }
                ],
                max_completion_tokens=4000
            )
            page_text = response.choices[0].message.content or ""
            text += f"\n\n--- –°—Ç—Ä–∞–Ω–∏—Ü–∞ {page_index + 1} ---\n\n{page_text}"

    logger.info(f"‚úÖ –ò–∑–≤–ª–µ—á–µ–Ω —Ç–µ–∫—Å—Ç –∏–∑ PDF {filename}, —Å–∏–º–≤–æ–ª–æ–≤: {len(text)}")

    if not text.strip():
        logger.error(f"‚ùå –ù–µ —É–¥–∞–ª–æ—Å—å –∏–∑–≤–ª–µ—á—å —Ç–µ–∫—Å—Ç –∏–∑ {filename}")
        return "[–î–æ–∫—É–º–µ–Ω—Ç –ø—É—Å—Ç –∏–ª–∏ –Ω–µ —É–¥–∞–ª–æ—Å—å —Ä–∞—Å–ø–æ–∑–Ω–∞—Ç—å —Ç–µ–∫—Å—Ç]"

    return text.strip()


async def extract_text_from_docx(content: bytes, filename: str) -> str:
    """–ò–∑–≤–ª–µ–∫–∞–µ—Ç —Ç–µ–∫—Å—Ç –∏–∑ DOCX, –≤–∫–ª—é—á–∞—è —Ç–∞–±–ª–∏—Ü—ã."""
    try:
        from docx import Document as DocxDocument
    except Exception:
        logger.error("python-docx –Ω–µ —É—Å—Ç–∞–Ω–æ–≤–ª–µ–Ω")
        return ""

    import io
    text_parts: List[str] = []
    docx_stream = io.BytesIO(content)
    doc = DocxDocument(docx_stream)

    # –ü–∞—Ä–∞–≥—Ä–∞—Ñ—ã
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            text_parts.append(p.text.strip())

    # –¢–∞–±–ª–∏—Ü—ã
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                text_parts.append(" | ".join(cells))

    result = "\n".join(text_parts)
    logger.info(f"‚úÖ –ò–∑–≤–ª–µ—á–µ–Ω —Ç–µ–∫—Å—Ç –∏–∑ DOCX {filename}, —Å–∏–º–≤–æ–ª–æ–≤: {len(result)}")
    return result


async def extract_text_from_any(content: bytes, filename: str) -> str:
    """–û–ø—Ä–µ–¥–µ–ª—è–µ—Ç —Ç–∏–ø (docx/pdf) –∏ –∏–∑–≤–ª–µ–∫–∞–µ—Ç —Ç–µ–∫—Å—Ç —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤—É—é—â–∏–º —Å–ø–æ—Å–æ–±–æ–º."""
    lower = (filename or "").lower()
    if lower.endswith('.docx'):
        return await extract_text_from_docx(content, filename)
    # default: PDF
    return await extract_text_from_pdf(content, filename)


@retry(stop=stop_after_attempt(RETRY_MAX_ATTEMPTS), wait=wait_exponential(multiplier=RETRY_WAIT_EXPONENTIAL_MULTIPLIER, min=4, max=RETRY_WAIT_EXPONENTIAL_MAX))
async def segment_requirements(tz_text: str) -> List[Dict[str, Any]]:
    """–°–µ–≥–º–µ–Ω—Ç–∏—Ä—É–µ—Ç –¢–ó –Ω–∞ –æ—Ç–¥–µ–ª—å–Ω—ã–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –∏—Å–ø–æ–ª—å–∑—É—è GPT."""
    # –§–æ—Ä–º–∏—Ä—É–µ–º –ø—Ä–æ–º–ø—Ç –∏–∑ –∑–∞–≥—Ä—É–∂–µ–Ω–Ω–æ–≥–æ —à–∞–±–ª–æ–Ω–∞ + —Ç–µ–∫—Å—Ç –¢–ó
    prompt = f"""{REQUIREMENTS_EXTRACTION_PROMPT}

–¢–µ–∫—Å—Ç –¢–ó:
{tz_text[:10000]}"""  # –û–≥—Ä–∞–Ω–∏—á–∏–≤–∞–µ–º –¥–æ 10000 —Å–∏–º–≤–æ–ª–æ–≤

    response = await client.chat.completions.create(
        model=OPENAI_MODEL,  # –ò—Å–ø–æ–ª—å–∑—É–µ–º –º–æ–¥–µ–ª—å –∏–∑ –∫–æ–Ω—Ñ–∏–≥–∞
        messages=[{"role": "user", "content": prompt}],
        response_format={"type": "json_object"}
    )

    try:
        response_text = response.choices[0].message.content
        logger.info(f"üìÑ GPT response preview: {response_text[:500]}...")
        
        data = json.loads(response_text)
        logger.info(f"üìä Parsed JSON keys: {list(data.keys())}")
        
        requirements = data.get("requirements", [])
        
        if not requirements:
            logger.warning(f"‚ö†Ô∏è –ü—É—Å—Ç–æ–π —Å–ø–∏—Å–æ–∫ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π! –ü–æ–ª–Ω—ã–π –æ—Ç–≤–µ—Ç GPT: {response_text}")
        
        logger.info(f"‚úÖ –ò–∑–≤–ª–µ—á–µ–Ω–æ {len(requirements)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π")
        return requirements
    except json.JSONDecodeError as e:
        logger.error(f"‚ùå Failed to parse requirements JSON: {e}")
        logger.error(f"‚ùå Response text: {response.choices[0].message.content}")
        raise ValueError("Failed to parse requirements JSON")




async def _get_file_size(file: UploadFile) -> int:
    """–ü–æ–ª—É—á–∞–µ—Ç —Ä–∞–∑–º–µ—Ä —Ñ–∞–π–ª–∞ –≤ –±–∞–π—Ç–∞—Ö."""
    content = await file.read()
    await file.seek(0)
    return len(content)


# ============================
# PYDANTIC –ú–û–î–ï–õ–ò
# ============================

class RequirementAnalysis(BaseModel):
    """–†–µ–∑—É–ª—å—Ç–∞—Ç –∞–Ω–∞–ª–∏–∑–∞ –æ–¥–Ω–æ–≥–æ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è"""
    number: int
    requirement: str
    status: str
    confidence: int
    solution_description: str
    reference: str
    evidence_text: Optional[str] = None  # –ö–æ–Ω–∫—Ä–µ—Ç–Ω—ã–π —Ç–µ–∫—Å—Ç —Å –ª–∏—Å—Ç–∞ –¥–ª—è –ø–æ–∏—Å–∫–∞
    discrepancies: str
    section: Optional[str] = None


class AnalysisResponse(BaseModel):
    """–û—Ç–≤–µ—Ç —Å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞–º–∏ –∞–Ω–∞–ª–∏–∑–∞"""
    stage: str
    req_type: str
    requirements: List[RequirementAnalysis]
    summary: str
    sheet_to_pdf_mapping: Optional[Dict[str, int]] = {}  # Mapping: sheet_number ‚Üí pdf_page_number


# ============================
# –ì–õ–û–ë–ê–õ–¨–ù–û–ï –°–û–°–¢–û–Ø–ù–ò–ï –ê–ù–ê–õ–ò–ó–ê
# ============================

analysis_status = {
    "current_stage": None,
    "progress": 0,
    "stage_name": "",
    "total_stages": 3,
    "start_time": None,
    "is_running": False
}

def update_analysis_status(stage_num: int, stage_name: str, progress: int):
    """–û–±–Ω–æ–≤–ª—è–µ—Ç –≥–ª–æ–±–∞–ª—å–Ω—ã–π —Å—Ç–∞—Ç—É—Å –∞–Ω–∞–ª–∏–∑–∞"""
    analysis_status.update({
        "current_stage": stage_num,
        "progress": progress,
        "stage_name": stage_name,
        "is_running": True
    })
    logger.info(f"üìä Status updated: Stage {stage_num}/3 - {stage_name} - {progress}%")

def reset_analysis_status():
    """–°–±—Ä–∞—Å—ã–≤–∞–µ—Ç —Å—Ç–∞—Ç—É—Å –∞–Ω–∞–ª–∏–∑–∞"""
    analysis_status.update({
        "current_stage": None,
        "progress": 0,
        "stage_name": "",
        "is_running": False,
        "start_time": None
    })

# After analysis_status definition
extraction_status = {
    "current_stage": None,
    "progress": 0,
    "stage_name": "",
    "total_stages": 2,
    "start_time": None,
    "is_running": False
}

def update_extraction_status(stage_num: int, stage_name: str, progress: int):
    extraction_status.update({
        "current_stage": stage_num,
        "progress": progress,
        "stage_name": stage_name,
        "is_running": True
    })
    logger.info(f"üìä Extraction status updated: Stage {stage_num}/2 - {stage_name} - {progress}%")

def reset_extraction_status():
    extraction_status.update({
        "current_stage": None,
        "progress": 0,
        "stage_name": "",
        "is_running": False,
        "start_time": None
    })

# ============================
# FASTAPI –ü–†–ò–õ–û–ñ–ï–ù–ò–ï
# ============================

app = FastAPI(
    title="Document Analysis API (Vision Mode)",
    description="API –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ —Å—Ç—Ä–æ–∏—Ç–µ–ª—å–Ω–æ–π –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏ —Å –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏–µ–º Vision API",
    version="6.0.0-vision"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============================
# API ENDPOINTS
# ============================

@app.get("/")
async def root():
    """Health check"""
    return {
        "status": "ok",
        "service": "Document Analysis API (VISION MODE)",
        "architecture": "Two-step: 1) Extract requirements from TZ, 2) Analyze project",
        "provider": "openai",
        "model": OPENAI_MODEL,
        "max_file_size_mb": MAX_FILE_SIZE_MB
    }


@app.get("/status")
async def get_analysis_status():
    """–ü–æ–ª—É—á–∏—Ç—å —Ç–µ–∫—É—â–∏–π —Å—Ç–∞—Ç—É—Å –∞–Ω–∞–ª–∏–∑–∞"""
    return analysis_status

@app.get("/extraction_status")
async def get_extraction_status():
    return extraction_status


@app.post("/extract_requirements")
async def extract_requirements_endpoint(
    request: Request,
    tz_document: UploadFile = File(...)
):
    """
    –®–∞–≥ 1: –ò–∑–≤–ª–µ—á–µ–Ω–∏–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π –∏–∑ –¢–ó.
    –ü–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ—Ç PDF (—Ç–µ–∫—Å—Ç –∏ –∏–∑–æ–±—Ä–∞–∂–µ–Ω–∏—è) –∏ DOCX.
    
    Returns:
        List of requirements with structure:
        [
            {
                "number": 1,
                "text": "requirement text",
                "section": "category",
                "trace_id": "req-1",
                "selected": true  # –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é –≤—Å–µ –≤—ã–±—Ä–∞–Ω—ã
            }
        ]
    """
    try:
        if await request.is_disconnected():
            logger.warning("‚ö†Ô∏è Client disconnected before extraction started.")
            raise HTTPException(status_code=499, detail="Client disconnected")
        
        # –°–±—Ä–∞—Å—ã–≤–∞–µ–º –∏ –∏–Ω–∏—Ü–∏–∞–ª–∏–∑–∏—Ä—É–µ–º —Å—Ç–∞—Ç—É—Å
        reset_extraction_status()
        update_extraction_status(1, "–ü–æ–¥–≥–æ—Ç–æ–≤–∫–∞ –¥–æ–∫—É–º–µ–Ω—Ç–∞", 5)
        
        logger.info(f"üìã [STEP 1] –ò–∑–≤–ª–µ—á–µ–Ω–∏–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π –∏–∑ {tz_document.filename}")
        
        # –ü—Ä–æ–≤–µ—Ä–∫–∞ —Ä–∞–∑–º–µ—Ä–∞ —Ñ–∞–π–ª–∞
        file_size = await _get_file_size(tz_document)
        if file_size > MAX_FILE_SIZE_BYTES:
            raise HTTPException(
                status_code=413,
                detail=f"–§–∞–π–ª {tz_document.filename} —Å–ª–∏—à–∫–æ–º –±–æ–ª—å—à–æ–π ({file_size / 1024 / 1024:.2f} MB). –ú–∞–∫—Å–∏–º—É–º: {MAX_FILE_SIZE_MB} MB"
            )
        
        # –ß–∏—Ç–∞–µ–º —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ
        await tz_document.seek(0)
        tz_content = await tz_document.read()
        
        logger.info(f"üìä File size: {len(tz_content) / 1024:.1f} KB")
        update_extraction_status(1, "–ò–∑–≤–ª–µ—á–µ–Ω–∏–µ —Ç–µ–∫—Å—Ç–∞ –∏–∑ –¥–æ–∫—É–º–µ–Ω—Ç–∞", 20)
        
        # –ò–∑–≤–ª–µ–∫–∞–µ–º —Ç–µ–∫—Å—Ç –∏–∑ –¥–æ–∫—É–º–µ–Ω—Ç–∞
        logger.info("üìÑ Extracting text from TZ document...")
        if await request.is_disconnected():
            logger.warning("‚ö†Ô∏è Client disconnected during text extraction")
            raise HTTPException(status_code=499, detail="Client disconnected")
        
        tz_text = await extract_text_from_any(tz_content, tz_document.filename)
        update_extraction_status(1, "–¢–µ–∫—Å—Ç —É—Å–ø–µ—à–Ω–æ –∏–∑–≤–ª–µ—á—ë–Ω", 50)
        
        # –°–µ–≥–º–µ–Ω—Ç–∏—Ä—É–µ–º —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è
        logger.info("‚úÇÔ∏è Segmenting requirements...")
        update_extraction_status(2, "–°–µ–≥–º–µ–Ω—Ç–∞—Ü–∏—è —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π", 60)
        
        if await request.is_disconnected():
            logger.warning("‚ö†Ô∏è Client disconnected during segmentation")
            raise HTTPException(status_code=499, detail="Client disconnected")
        
        requirements = await segment_requirements(tz_text)
        update_extraction_status(2, "–°–µ–≥–º–µ–Ω—Ç–∞—Ü–∏—è –∑–∞–≤–µ—Ä—à–µ–Ω–∞", 90)
        
        if not requirements:
            raise HTTPException(status_code=400, detail="No requirements extracted from TZ")
        
        # –î–æ–±–∞–≤–ª—è–µ–º –ø–æ–ª–µ selected=true –¥–ª—è –≤—Å–µ—Ö —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é
        for req in requirements:
            req['selected'] = True
        
        logger.info(f"‚úÖ Successfully extracted {len(requirements)} requirements")
        update_extraction_status(2, f"–ò–∑–≤–ª–µ—á–µ–Ω–æ {len(requirements)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π", 100)
        
        return {
            "success": True,
            "total_requirements": len(requirements),
            "requirements": requirements
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"‚ùå Error extracting requirements: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"–û—à–∏–±–∫–∞ –∏–∑–≤–ª–µ—á–µ–Ω–∏—è —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π: {str(e)}")


@app.post("/analyze", response_model=AnalysisResponse)
async def analyze_documentation(
    request: Request,
    stage: str = Form(...),
    requirements_json: str = Form(...),  # JSON string —Å —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è–º–∏ –∏–∑ —à–∞–≥–∞ 1
    doc_document: UploadFile = File(...)
):
    """
    –®–∞–≥ 2: –ê–Ω–∞–ª–∏–∑ –ø—Ä–æ–µ–∫—Ç–Ω–æ–π –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏ –ø–æ –≥–æ—Ç–æ–≤—ã–º —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è–º.
    
    Args:
        stage: –°—Ç–∞–¥–∏—è –ø—Ä–æ–µ–∫—Ç–∞ (–ì–ö, –§–≠, –≠–ü)
        requirements_json: JSON string —Å —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è–º–∏ –∏–∑ —à–∞–≥–∞ 1
        doc_document: –§–∞–π–ª –ø—Ä–æ–µ–∫—Ç–Ω–æ–π –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏ (PDF)
    """
    try:
        # –ü—Ä–æ–≤–µ—Ä—è–µ–º, –Ω–µ –æ—Ç–∫–ª—é—á–∏–ª—Å—è –ª–∏ –∫–ª–∏–µ–Ω—Ç
        if await request.is_disconnected():
            logger.warning("‚ö†Ô∏è Client disconnected before analysis started. Aborting.")
            raise HTTPException(status_code=499, detail="Client disconnected")

        logger.info(f"üìã [STEP 2] –ê–Ω–∞–ª–∏–∑ –ø—Ä–æ–µ–∫—Ç–Ω–æ–π –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏. –°—Ç–∞–¥–∏—è: {stage}")

        # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å –∏ –Ω–∞—á–∏–Ω–∞–µ–º –∞–Ω–∞–ª–∏–∑
        reset_analysis_status()

        # ============================================================
        # –≠–¢–ê–ü 1: –ü–∞—Ä—Å–∏–Ω–≥ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π –∏–∑ JSON
        # ============================================================

        try:
            requirements = json.loads(requirements_json)
            logger.info(f"üìã –ü–æ–ª—É—á–µ–Ω–æ {len(requirements)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π –∏–∑ —à–∞–≥–∞ 1")
            update_analysis_status(1, "–ü–æ–¥–≥–æ—Ç–æ–≤–∫–∞ –¥–∞–Ω–Ω—ã—Ö", 5)
            
            # –§–∏–ª—å—Ç—Ä—É–µ–º —Ç–æ–ª—å–∫–æ –≤—ã–±—Ä–∞–Ω–Ω—ã–µ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è (selected=true)
            selected_requirements = [req for req in requirements if req.get('selected', True)]
            logger.info(f"‚úÖ –í—ã–±—Ä–∞–Ω–æ {len(selected_requirements)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞")
            
            if not selected_requirements:
                raise HTTPException(status_code=400, detail="No requirements selected for analysis")
            
            requirements = selected_requirements
            
        except json.JSONDecodeError as e:
            logger.error(f"‚ùå Failed to parse requirements JSON: {e}")
            raise HTTPException(status_code=400, detail="Invalid requirements JSON format")

        # ============================================================
        # –≠–¢–ê–ü 2: –ü—Ä–æ–≤–µ—Ä–∫–∞ —Ñ–∞–π–ª–∞ –ø—Ä–æ–µ–∫—Ç–Ω–æ–π –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏
        # ============================================================
        
        # –ü—Ä–æ–≤–µ—Ä–∫–∞ —Ä–∞–∑–º–µ—Ä–∞ —Ñ–∞–π–ª–∞
        file_size = await _get_file_size(doc_document)
        if file_size > MAX_FILE_SIZE_BYTES:
            raise HTTPException(
                status_code=413,
                detail=f"–§–∞–π–ª {doc_document.filename} —Å–ª–∏—à–∫–æ–º –±–æ–ª—å—à–æ–π ({file_size / 1024 / 1024:.2f} MB). –ú–∞–∫—Å–∏–º—É–º: {MAX_FILE_SIZE_MB} MB"
            )

        # Read contents
        await doc_document.seek(0)
        doc_content = await doc_document.read()

        logger.info(f"üìä File size - DOC: {len(doc_content) / 1024:.1f} KB")

        # ============================================================
        # –≠–¢–ê–ü 3 [STAGE 1]: –ò–∑–≤–ª–µ—á–µ–Ω–∏–µ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö —Å—Ç—Ä–∞–Ω–∏—Ü
        # ============================================================

        logger.info("üìã [STEP 1/3] STAGE 1: Extracting page metadata...")
        pages_metadata = await extract_page_metadata(doc_content, doc_document.filename, max_pages=150)

        # üîß –ö–†–ò–¢–ò–ß–ï–°–ö–ò –í–ê–ñ–ù–û: –ù–æ—Ä–º–∞–ª–∏–∑—É–µ–º –Ω–æ–º–µ—Ä–∞ –ª–∏—Å—Ç–æ–≤ –∫ –¢–û–õ–¨–ö–û –¶–ò–§–†–û–í–û–ú–£ —Ñ–æ—Ä–º–∞—Ç—É
        logger.info("üìã [STAGE 1] –ù–æ—Ä–º–∞–ª–∏–∑–∞—Ü–∏—è –Ω–æ–º–µ—Ä–æ–≤ –ª–∏—Å—Ç–æ–≤ –∫ —Ü–∏—Ñ—Ä–æ–≤–æ–º—É —Ñ–æ—Ä–º–∞—Ç—É...")
        for page_meta in pages_metadata:
            original_sheet_num = page_meta.get('sheet_number')
            if original_sheet_num and original_sheet_num != "N/A":
                normalized_sheet_num = normalize_sheet_number_to_digit(original_sheet_num)
                page_meta['sheet_number'] = normalized_sheet_num
                if normalized_sheet_num != original_sheet_num:
                    logger.info(f"üìã –ù–æ—Ä–º–∞–ª–∏–∑–æ–≤–∞–Ω–æ: —Å—Ç—Ä.{page_meta.get('page')} '{original_sheet_num}' ‚Üí '{normalized_sheet_num}'")

        logger.info(f"‚úÖ [STAGE 1] –ù–æ—Ä–º–∞–ª–∏–∑–∞—Ü–∏—è –∑–∞–≤–µ—Ä—à–µ–Ω–∞. –í—Å–µ –Ω–æ–º–µ—Ä–∞ –ª–∏—Å—Ç–æ–≤ —Ç–µ–ø–µ—Ä—å –≤ —Ü–∏—Ñ—Ä–æ–≤–æ–º —Ñ–æ—Ä–º–∞—Ç–µ.")
        update_analysis_status(1, "–ò–∑–≤–ª–µ—á–µ–Ω–∏–µ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö", 33)

        # –°–æ–∑–¥–∞–µ–º mapping: sheet_number ‚Üí pdf_page_number –¥–ª—è –Ω–∞–≤–∏–≥–∞—Ü–∏–∏
        sheet_to_pdf_mapping = {}
        for page_meta in pages_metadata:
            pdf_page = page_meta.get('page')
            sheet_num = page_meta.get('sheet_number', str(pdf_page))

            # –ü—Ä–æ–ø—É—Å–∫–∞–µ–º –Ω–µ–≤–∞–ª–∏–¥–Ω—ã–µ –Ω–æ–º–µ—Ä–∞
            if not sheet_num or sheet_num == "N/A":
                continue

            # –ù–æ—Ä–º–∞–ª–∏–∑—É–µ–º –Ω–æ–º–µ—Ä –ª–∏—Å—Ç–∞ –¥–ª—è –Ω–∞–¥–µ–∂–Ω–æ–≥–æ –ø–æ–∏—Å–∫–∞
            sheet_num_normalized = str(sheet_num).strip()

            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –¥—É–±–ª–∏–∫–∞—Ç—ã –∏ –ø—Ä–µ–¥—É–ø—Ä–µ–∂–¥–∞–µ–º
            if sheet_num_normalized in sheet_to_pdf_mapping:
                logger.warning(f"‚ö†Ô∏è [STAGE 1] –î—É–±–ª–∏–∫–∞—Ç –Ω–æ–º–µ—Ä–∞ –ª–∏—Å—Ç–∞ '{sheet_num_normalized}': PDF —Å—Ç—Ä–∞–Ω–∏—Ü—ã {sheet_to_pdf_mapping[sheet_num_normalized]} –∏ {pdf_page}")
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –ø–µ—Ä–≤–æ–µ –≤—Ö–æ–∂–¥–µ–Ω–∏–µ (–æ–±—ã—á–Ω–æ –æ–Ω–æ –∫–æ—Ä—Ä–µ–∫—Ç–Ω–æ–µ)
                continue

            sheet_to_pdf_mapping[sheet_num_normalized] = pdf_page

            # –î–æ–±–∞–≤–ª—è–µ–º –∞–ª—å—Ç–µ—Ä–Ω–∞—Ç–∏–≤–Ω—ã–µ –≤–∞—Ä–∏–∞–Ω—Ç—ã –∑–∞–ø–∏—Å–∏ –¥–ª—è –Ω–∞–¥–µ–∂–Ω–æ—Å—Ç–∏
            # –ù–∞–ø—Ä–∏–º–µ—Ä: "–ê–†-01" ‚Üí —Ç–∞–∫–∂–µ –¥–æ—Å—Ç—É–ø–µ–Ω –∫–∞–∫ "–ê–†-1", "–∞—Ä-01", "–ê–†01"
            alternatives = []

            # –í–∞—Ä–∏–∞–Ω—Ç –±–µ–∑ –Ω—É–ª–µ–π: "–ê–†-01" ‚Üí "–ê–†-1"
            if '-' in sheet_num_normalized:
                parts = sheet_num_normalized.split('-')
                if len(parts) == 2 and parts[1].isdigit():
                    alternatives.append(f"{parts[0]}-{int(parts[1])}")

            # –í–∞—Ä–∏–∞–Ω—Ç –±–µ–∑ –¥–µ—Ñ–∏—Å–∞: "–ê–†-01" ‚Üí "–ê–†01"
            alternatives.append(sheet_num_normalized.replace('-', ''))
            alternatives.append(sheet_num_normalized.replace('‚Äì', ''))  # em-dash
            alternatives.append(sheet_num_normalized.replace('‚Äî', ''))  # en-dash

            # –í–∞—Ä–∏–∞–Ω—Ç –≤ –Ω–∏–∂–Ω–µ–º —Ä–µ–≥–∏—Å—Ç—Ä–µ
            alternatives.append(sheet_num_normalized.lower())

            # –î–æ–±–∞–≤–ª—è–µ–º –≤—Å–µ –∞–ª—å—Ç–µ—Ä–Ω–∞—Ç–∏–≤—ã
            for alt in alternatives:
                if alt and alt != sheet_num_normalized and alt not in sheet_to_pdf_mapping:
                    sheet_to_pdf_mapping[alt] = pdf_page

        logger.info(f"üìä [STAGE 1] –°–æ–∑–¥–∞–Ω mapping –ª–∏—Å—Ç–æ–≤: {len(sheet_to_pdf_mapping)} –≤–∞—Ä–∏–∞–Ω—Ç–æ–≤ –¥–ª—è {len(pages_metadata)} —Å—Ç—Ä–∞–Ω–∏—Ü")
        logger.info(f"üìä [STAGE 1] –ü—Ä–∏–º–µ—Ä—ã: {list(sheet_to_pdf_mapping.items())[:10]}...")


        # ============================================================
        # –≠–¢–ê–ü 2 [STAGE 2]: –ö–æ–Ω–≤–µ—Ä—Ç–∞—Ü–∏—è –≤ –Ω–∏–∑–∫–æ–µ –∫–∞—á–µ—Å—Ç–≤–æ –∏ –æ—Ü–µ–Ω–∫–∞ —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω–æ—Å—Ç–∏
        # ============================================================

        logger.info("üì§ [STEP 2/3] STAGE 2: Converting to low-res and assessing relevance...")
        update_analysis_status(2, "–ü—Ä–µ—Ñ–∏–ª—å—Ç—Ä —Å—Ç—Ä–∞–Ω–∏—Ü –ø–æ —Ç–µ–∫—Å—Ç—É", 40)

        # –¢–µ–∫—Å—Ç–æ–≤—ã–π –ø—Ä–µ—Ñ–∏–ª—å—Ç—Ä —Å—Ç—Ä–∞–Ω–∏—Ü
        page_texts_quick = _extract_page_texts_quick(doc_content, max_pages=STAGE2_MAX_PAGES)
        candidate_pages = _simple_candidate_pages(requirements, page_texts_quick, per_req=7, cap_total=30)
        logger.info(f"üìÑ [STAGE 2] –¢–µ–∫—Å—Ç–æ–≤—ã–π –ø—Ä–µ—Ñ–∏–ª—å—Ç—Ä –≤—ã–±—Ä–∞–ª —Å—Ç—Ä–∞–Ω–∏—Ü—ã: {candidate_pages[:10]}{'...' if len(candidate_pages) > 10 else ''}")

        update_analysis_status(2, "–ò–∑–≤–ª–µ—á–µ–Ω–∏–µ –≤—ã–±—Ä–∞–Ω–Ω—ã—Ö —Å—Ç—Ä–∞–Ω–∏—Ü", 50)

        # –ò–∑–≤–ª–µ–∫–∞–µ–º —Ç–æ–ª—å–∫–æ –≤—ã–±—Ä–∞–Ω–Ω—ã–µ —Å—Ç—Ä–∞–Ω–∏—Ü—ã –≤ low-res
        doc_images_low, page_numbers_kept = await extract_selected_pdf_pages_as_images(
            doc_content, doc_document.filename, selected_pages=candidate_pages,
            detail=STAGE2_DETAIL, dpi=STAGE2_DPI, quality=STAGE2_QUALITY
        )

        update_analysis_status(2, "–û—Ü–µ–Ω–∫–∞ —Ä–µ–ª–µ–≤–∞–Ω—Ç–Ω–æ—Å—Ç–∏ —Å—Ç—Ä–∞–Ω–∏—Ü", 60)

        page_mapping = await assess_page_relevance(pages_metadata, doc_images_low, requirements, page_numbers=page_numbers_kept)
        update_analysis_status(2, "–†–µ–ª–µ–≤–∞–Ω—Ç–Ω–æ—Å—Ç—å –æ–ø—Ä–µ–¥–µ–ª–µ–Ω–∞", 66)

        # ============================================================
        # –≠–¢–ê–ü 5: –ü–æ–¥–≥–æ—Ç–æ–≤–∫–∞ system prompt
        # ============================================================

        system_prompt = get_analysis_system_prompt(stage, "–¢–ó")

        # ============================================================
        # –≠–¢–ê–ü 6 [STAGE 3]: –ì—Ä—É–ø–ø–∏—Ä–æ–≤–∫–∞ –∏ –∞–Ω–∞–ª–∏–∑ —Å –≤—ã—Å–æ–∫–∏–º —Ä–∞–∑—Ä–µ—à–µ–Ω–∏–µ–º
        # ============================================================

        logger.info(f"üîç [STEP 3/3] STAGE 3: Analyzing with high-resolution images...")
        update_analysis_status(3, "–î–µ—Ç–∞–ª—å–Ω—ã–π –∞–Ω–∞–ª–∏–∑ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π", 70)
        analyzed_reqs = []

        # –ì—Ä—É–ø–ø–∏—Ä—É–µ–º —Ç—Ä–µ–±–æ–≤–∞–Ω–∏—è –ø–æ –æ–±—â–∏–º —Å—Ç—Ä–∞–Ω–∏—Ü–∞–º –¥–ª—è –æ–ø—Ç–∏–º–∏–∑–∞—Ü–∏–∏
        from collections import defaultdict
        page_to_reqs = defaultdict(list)

        for req in requirements:
            req_pages = page_mapping.get(req['number'], [])
            if not req_pages:  # Fallback - –ø–µ—Ä–≤—ã–µ 20 —Å—Ç—Ä–∞–Ω–∏—Ü
                req_pages = page_numbers_kept[:min(20, len(page_numbers_kept))] if page_numbers_kept else list(range(1, min(21, len(doc_images_low) + 1)))

            pages_key = tuple(sorted(req_pages))
            page_to_reqs[pages_key].append(req)

        logger.info(f"üì¶ [STAGE 3] –°–æ–∑–¥–∞–Ω–æ {len(page_to_reqs)} –≥—Ä—É–ø–ø –ø–æ –æ–±—â–∏–º —Å—Ç—Ä–∞–Ω–∏—Ü–∞–º")
        total_groups = len(page_to_reqs)

        for group_idx, (pages_key, reqs_group) in enumerate(page_to_reqs.items(), 1):
            if await request.is_disconnected():
                logger.warning(f"‚ö†Ô∏è Client disconnected at group {group_idx}/{total_groups}")
                # –í–æ–∑–≤—Ä–∞—â–∞–µ–º —á–∞—Å—Ç–∏—á–Ω—ã–µ —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã
                return AnalysisResponse(
                    stage=stage,
                    req_type="–¢–ó",
                    requirements=analyzed_reqs,
                    summary=f"–ê–Ω–∞–ª–∏–∑ –ø—Ä–µ—Ä–≤–∞–Ω: –∫–ª–∏–µ–Ω—Ç –æ—Ç–∫–ª—é—á–∏–ª—Å—è –ø–æ—Å–ª–µ –æ–±—Ä–∞–±–æ—Ç–∫–∏ {len(analyzed_reqs)}/{len(requirements)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π (–≥—Ä—É–ø–ø–∞ {group_idx}/{total_groups})"
                )

            # –î–∏–Ω–∞–º–∏—á–µ—Å–∫–∏–π –ø—Ä–æ–≥—Ä–µ—Å—Å: 70% + (0-25% –≤ –∑–∞–≤–∏—Å–∏–º–æ—Å—Ç–∏ –æ—Ç –≥—Ä—É–ø–ø—ã)
            progress = 70 + int((group_idx / total_groups) * 25)
            update_analysis_status(
                3, 
                f"–ê–Ω–∞–ª–∏–∑ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π ({group_idx}/{total_groups} –≥—Ä—É–ø–ø)", 
                progress
            )

            logger.info(f"üì¶ [STAGE 3] [{group_idx}/{total_groups}] Analyzing {len(reqs_group)} requirements on {len(pages_key)} pages")

            # –†–∞–∑–±–∏–≤–∞–µ–º –Ω–∞ –ø–∞–∫–µ—Ç—ã –ø–æ N —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π –µ—Å–ª–∏ –≥—Ä—É–ø–ø–∞ –±–æ–ª—å—à–∞—è
            for batch_start in range(0, len(reqs_group), STAGE3_BATCH_SIZE):
                batch = reqs_group[batch_start:batch_start + STAGE3_BATCH_SIZE]

                batch_results = await analyze_batch_with_high_detail(
                    system_prompt=system_prompt,
                    doc_content=doc_content,
                    page_numbers=list(pages_key),
                    requirements_batch=batch,
                    request=request,
                    pages_metadata=pages_metadata
                )

                if not batch_results:
                    # –ö–ª–∏–µ–Ω—Ç –æ—Ç–∫–ª—é—á–∏–ª—Å—è –≤–æ –≤—Ä–µ–º—è batch –∞–Ω–∞–ª–∏–∑–∞
                    return AnalysisResponse(
                        stage=stage,
                        req_type="–¢–ó",
                        requirements=analyzed_reqs,
                        summary=f"–ê–Ω–∞–ª–∏–∑ –ø—Ä–µ—Ä–≤–∞–Ω: –∫–ª–∏–µ–Ω—Ç –æ—Ç–∫–ª—é—á–∏–ª—Å—è –≤–æ –≤—Ä–µ–º—è batch –∞–Ω–∞–ª–∏–∑–∞. –û–±—Ä–∞–±–æ—Ç–∞–Ω–æ {len(analyzed_reqs)}/{len(requirements)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π"
                    )

                analyzed_reqs.extend(batch_results)

        # –°–æ—Ä—Ç–∏—Ä—É–µ–º –ø–æ –∏—Å—Ö–æ–¥–Ω–æ–º—É –ø–æ—Ä—è–¥–∫—É –∏–∑ –¢–ó
        analyzed_reqs.sort(key=lambda r: r.number)

        # ============================================================
        # –≠–¢–ê–ü 7: –ì–µ–Ω–µ—Ä–∞—Ü–∏—è —Å–≤–æ–¥–∫–∏
        # ============================================================

        logger.info("üìù [STEP 3/3] Generating summary...")
        update_analysis_status(3, "–ì–µ–Ω–µ—Ä–∞—Ü–∏—è –æ—Ç—á–µ—Ç–∞", 95)

        if await request.is_disconnected():
            logger.warning("‚ö†Ô∏è Client disconnected before summary")
            # –í–æ–∑–≤—Ä–∞—â–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç—ã –±–µ–∑ summary
            return AnalysisResponse(
                stage=stage,
                req_type="–¢–ó",
                requirements=analyzed_reqs,
                summary=f"–ê–Ω–∞–ª–∏–∑ –∑–∞–≤–µ—Ä—à–µ–Ω, –Ω–æ –∫–ª–∏–µ–Ω—Ç –æ—Ç–∫–ª—é—á–∏–ª—Å—è –ø–µ—Ä–µ–¥ –≥–µ–Ω–µ—Ä–∞—Ü–∏–µ–π —Å–≤–æ–¥–∫–∏. –ü—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–æ {len(analyzed_reqs)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π."
            )

        # –°—Ç–∞—Ç–∏—Å—Ç–∏–∫–∞ –¥–ª—è —Å–≤–æ–¥–∫–∏
        total = len(analyzed_reqs)
        completed = sum(1 for r in analyzed_reqs if r.status == "–ü–æ–ª–Ω–æ—Å—Ç—å—é –∏—Å–ø–æ–ª–Ω–µ–Ω–æ")
        partial = sum(1 for r in analyzed_reqs if r.status == "–ß–∞—Å—Ç–∏—á–Ω–æ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ")
        not_done = sum(1 for r in analyzed_reqs if r.status == "–ù–µ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ")
        unclear = sum(1 for r in analyzed_reqs if r.status == "–¢—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è")

        summary = f"""–ê–Ω–∞–ª–∏–∑ –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏ –∑–∞–≤–µ—Ä—à–µ–Ω.

–ü—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–æ —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π: {total}
- –ü–æ–ª–Ω–æ—Å—Ç—å—é –∏—Å–ø–æ–ª–Ω–µ–Ω–æ: {completed} ({completed/total*100:.1f}%)
- –ß–∞—Å—Ç–∏—á–Ω–æ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ: {partial} ({partial/total*100:.1f}%)
- –ù–µ –∏—Å–ø–æ–ª–Ω–µ–Ω–æ: {not_done} ({not_done/total*100:.1f}%)
- –¢—Ä–µ–±—É–µ—Ç —É—Ç–æ—á–Ω–µ–Ω–∏—è: {unclear} ({unclear/total*100:.1f}%)

–°—Ä–µ–¥–Ω—è—è –¥–æ—Å—Ç–æ–≤–µ—Ä–Ω–æ—Å—Ç—å: {sum(r.confidence for r in analyzed_reqs)/total:.1f}%"""

        # ============================================================
        # –≠–¢–ê–ü 8 (–æ–ø—Ü–∏–æ–Ω–∞–ª—å–Ω–æ): –ü–æ–∏—Å–∫ –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏–π
        # ============================================================

        if STAGE4_ENABLED:
            logger.info("üîç STAGE 4: –ü–æ–∏—Å–∫ –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏–π –≤ –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ü–∏–∏...")
            try:
                contradictions_report = await find_contradictions(
                    pages_metadata=pages_metadata,
                    doc_content=doc_content,
                    requirements=requirements,
                    analyzed_reqs=analyzed_reqs
                )

                # –î–æ–±–∞–≤–ª—è–µ–º –æ—Ç—á–µ—Ç –æ –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏—è—Ö –∫ summary
                summary += "\n\n" + "="*80 + "\n\n" + contradictions_report

                logger.info("‚úÖ [STAGE 4] –ê–Ω–∞–ª–∏–∑ –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏–π –∑–∞–≤–µ—Ä—à–µ–Ω")

            except Exception as e:
                logger.error(f"‚ùå [STAGE 4] –û—à–∏–±–∫–∞ –ø—Ä–∏ –ø–æ–∏—Å–∫–µ –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏–π: {e}")
                summary += f"\n\n‚ö†Ô∏è –ê–Ω–∞–ª–∏–∑ –ø—Ä–æ—Ç–∏–≤–æ—Ä–µ—á–∏–π –Ω–µ –≤—ã–ø–æ–ª–Ω–µ–Ω: {str(e)}"
        else:
            logger.info("‚è≠Ô∏è [STAGE 4] –ü—Ä–æ–ø—É—â–µ–Ω (STAGE4_ENABLED=False)")

        # ============================================================
        # –í–æ–∑–≤—Ä–∞—Ç —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞
        # ============================================================

        # –í–æ–∑–≤—Ä–∞—â–∞–µ–º —Ä–µ–∑—É–ª—å—Ç–∞—Ç —Å mapping –ª–∏—Å—Ç–æ–≤
        parsed_result = AnalysisResponse(
            stage=stage,
            req_type="–¢–ó",
            requirements=analyzed_reqs,
            summary=summary,
            sheet_to_pdf_mapping=sheet_to_pdf_mapping  # –ù–æ–≤–æ–µ –ø–æ–ª–µ –¥–ª—è –Ω–∞–≤–∏–≥–∞—Ü–∏–∏
        )

        update_analysis_status(3, "–ê–Ω–∞–ª–∏–∑ –∑–∞–≤–µ—Ä—à–µ–Ω", 100)
        logger.info(f"‚úÖ [STEP 2] –ê–Ω–∞–ª–∏–∑ –∑–∞–≤–µ—Ä—à–µ–Ω —É—Å–ø–µ—à–Ω–æ. –ü—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–æ {len(analyzed_reqs)} —Ç—Ä–µ–±–æ–≤–∞–Ω–∏–π.")
        return parsed_result

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"‚ùå [STEP 2] –û—à–∏–±–∫–∞ –ø—Ä–∏ –∞–Ω–∞–ª–∏–∑–µ: {e}", exc_info=True)
        reset_analysis_status()  # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Å—Ç–∞—Ç—É—Å –ø—Ä–∏ –æ—à–∏–±–∫–µ
        raise HTTPException(status_code=500, detail=f"–û—à–∏–±–∫–∞ –∞–Ω–∞–ª–∏–∑–∞: {str(e)}")


# ============================
# –ó–ê–ü–£–°–ö –°–ï–†–í–ï–†–ê
# ============================

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000, log_level="info")
